from io import BytesIO
from openai import OpenAI
from app.prompts.document_reader_prompt import DocumentReaderPrompt
from langchain_openai import ChatOpenAI
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from app.models import Case, CaseBenefit
from datetime import datetime
from copy import deepcopy
from docxcompose.composer import Composer

_ = load_dotenv()


class AgentDocumentGenerator:
    """
    Agente especializado para geração de petições FAP usando modelos DOCX
    Processa templates Word e preenche com dados do banco de dados
    """
    model = None
    prompt = None

    def __init__(self, model_name="gpt-4o"):
        self.model = ChatOpenAI(model=model_name)
        self.prompt = DocumentReaderPrompt()

    def generate(self, file_id):
        """Gera resumo de documento usando IA"""
        result = self.model.invoke(
            [
                {"role": "system", "content": self.prompt.prompt_template()},
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": "Analise o documento e gere um resumo objetivo.",
                        },
                        {"type": "file", "file_id": file_id},
                    ],
                },
            ]
        )
        return result.content
    
    def generate_fap_petition(self, case_id, template_path=None):
        """
        Gera petição FAP preenchendo template DOCX com dados do banco
        Seleciona automaticamente o template correto baseado no fap_reason
        
        Args:
            case_id: ID do caso no banco de dados
            template_path: Caminho para o template DOCX (opcional, se não fornecido, seleciona automaticamente)
            
        Returns:
            Document: Documento DOCX preenchido
        """
        # Buscar dados do caso
        case = Case.query.get(case_id)
        if not case:
            raise ValueError(f"Caso {case_id} não encontrado")
        
        # Buscar benefícios do caso
        benefits = CaseBenefit.query.filter_by(case_id=case_id).all()
        
        # Selecionar template baseado no fap_reason se não foi especificado
        if template_path is None:
            template_path = self._select_template_by_fap_reason(case.fap_reason)
        
        # Carregar templates
        document_base = Document("templates_docx/modelo_acidente_trajeto_inicio.docx")
        document_content = Document(template_path)
        
        # Preencher campos do template BASE com dados do caso (cabeçalho/rodapé inicial)
        self._replace_placeholders_in_document(document_base, case, benefits)
        
        # Preencher campos do template CONTEÚDO com dados do caso
        self._replace_placeholders_in_document(document_content, case, benefits)
        
        # Adicionar benefícios nas tabelas
        self._add_benefits_to_tables(document_content, case, benefits)
        
        # Usar Composer para mesclar documentos preservando estilos e formatação
        composer = Composer(document_base)
        composer.append(document_content)

        return composer.doc
    
    def _append_document_content(self, base_doc, source_doc):
        """
        DEPRECATED: Método antigo substituído por docxcompose.Composer
        Mantido para referência mas não é mais usado.
        """
        pass
    
    def _select_template_by_fap_reason(self, fap_reason):
        """
        Seleciona o template correto baseado no motivo FAP
        
        Args:
            fap_reason: Código do motivo FAP
            
        Returns:
            str: Caminho para o template DOCX apropriado
        """
        # Mapeamento de motivos para templates
        template_mapping = {
            'inclusao_indevida_trajeto': 'templates_docx/modelo_acidente_trajeto.docx',
            'erro_material_cat': 'templates_docx/modelo_acidente_trajeto_erro_material.docx',
            'cat_trajeto_extemporanea': 'templates_docx/modelo_acidente_trajeto_extemporanea.docx'
        }
        
        # Retornar template específico ou padrão
        return template_mapping.get(fap_reason, 'templates_docx/modelo_acidente_trajeto.docx')
    
    def get_placeholders_preview(self, case, benefits):
        """
        Retorna um dicionário com os placeholders e seus valores para prévia
        Usado para exibir no HTML antes de gerar o documento
        
        Args:
            case: Objeto Case do banco de dados
            benefits: Lista de CaseBenefit
            
        Returns:
            dict: Dicionário com as categorias e placeholders
        """
        return {
            'Dados do Cliente': {
                '{{cliente_nome}}': case.client.name if case.client else 'Não informado',
                '{{cliente_cnpj}}': case.client.cnpj if case.client else 'Não informado',
                '{{cliente_endereco}}': self._format_address(case.client) if case.client else 'Não informado',
                '{{cliente_cidade}}': case.client.city if case.client else 'Não informado',
                '{{cliente_estado}}': case.client.state if case.client else 'Não informado',
            },
            'Dados do Caso': {
                '{{caso_titulo}}': case.title,
                '{{caso_tipo}}': case.case_type,
                '{{caso_numero}}': str(case.id),
            },
            'Dados FAP': {
                '{{fap_motivo}}': self._get_fap_reason_text(case.fap_reason) if case.fap_reason else 'Não informado',
                '{{ano_inicial_fap}}': str(case.fap_start_year) if case.fap_start_year else 'Não informado',
                '{{ano_final_fap}}': str(case.fap_end_year) if case.fap_end_year else 'Não informado',
                '{{anos_fap}}': self._format_years_range(case.fap_start_year, case.fap_end_year) or 'Não informado',
            },
            'Dados de Benefícios': {
                '{{total_beneficios}}': str(len(benefits)),
                '{{lista_beneficios}}': f"{len(benefits)} benefício(s) será(ão) listado(s)",
                '{{quantidade_acidentes}}': str(len(benefits)),
                '{{quantidade_acidentes_extenso}}': self._number_to_words(len(benefits)),
                '{{nome_segurado_exemplo}}': benefits[0].insured_name if benefits else 'Não informado',
                '{{nit_segurado_exemplo}}': benefits[0].insured_nit if benefits else 'Não informado',
                '{{data_acidente_segurado_exemplo}}': benefits[0].accident_date.strftime('%d/%m/%Y') if benefits and benefits[0].accident_date else 'Não informado',
                '{{data_inicio_beneficio_segurado_exemplo}}': benefits[0].data_inicio_beneficio.strftime('%d/%m/%Y') if benefits and benefits[0].data_inicio_beneficio else 'Não informado',
                '{{data_fim_beneficio_segurado_exemplo}}': benefits[0].data_fim_beneficio.strftime('%d/%m/%Y') if benefits and benefits[0].data_fim_beneficio else 'Não informado',
                '{{numero_beneficio}}': benefits[0].benefit_number if benefits else 'Não informado',
                '{{lesao_segurado_exemplo}}': 'Não informado',  # Campo adicional
                '{{resumo_evento_segurado_exemplo}}': 'Não informado',  # Campo adicional
                '{{data_afastamento_trabalho}}': 'Não informado',  # Campo adicional
                '{{data_laudo_medico_segurado_exemplo}}': 'Não informado',  # Campo adicional
                '{{numero_cat}}': benefits[0].numero_cat if benefits and benefits[0].numero_cat else 'Não informado',
                '{{numero_bo}}': benefits[0].numero_bo if benefits and benefits[0].numero_bo else 'Não informado',
            },
            'Valores': {
                '{{valor_causa}}': self._format_currency(case.value_cause) if case.value_cause else 'Não informado',
                '{{valor_causa_extenso}}': self._value_to_words(case.value_cause) if case.value_cause else 'Não informado',
            },
            'Datas': {
                '{{data_ajuizamento}}': case.filing_date.strftime('%d/%m/%Y') if case.filing_date else 'Não informado',
                '{{data_atual}}': datetime.now().strftime('%d/%m/%Y'),
                '{{mes_ano_atual}}': datetime.now().strftime('%B de %Y'),
                '{{vigencia_fap}}': self._format_years_range(case.fap_start_year, case.fap_end_year) or 'Não informado',
            },
            'Imagens': {
                '{{Imagem_cat}}': '[Imagem CAT será inserida]',
                '{{imagem_fap}}': '[Imagem FAP será inserida]',
                '{{imagem_info_beneficiario}}': '[Imagem Info Beneficiário será inserida]',
                '{{imagem_declaracao_beneficio}}': '[Imagem Declaração Benefício será inserida]',
                '{{imagem_inss_beneficiario}}': '[Imagem INSS Beneficiário será inserida]',
                '{{imagem_vigencia_beneficio}}': '[Imagem Vigência Benefício será inserida]',
            },
            'Dados da Vara': {
                '{{vara_nome}}': case.court.vara_name if case.court else 'Não informado',
                '{{vara_cidade}}': case.court.city if case.court else 'Não informado',
                '{{vara_estado}}': case.court.state if case.court else 'Não informado',
                '{{vara_completo}}': self._format_court(case.court) if case.court else 'Não informado',
            },
            'Resumos': {
                '{{fatos_resumo}}': case.facts_summary[:100] + '...' if case.facts_summary and len(case.facts_summary) > 100 else case.facts_summary or 'Não informado',
                '{{teses_resumo}}': case.thesis_summary[:100] + '...' if case.thesis_summary and len(case.thesis_summary) > 100 else case.thesis_summary or 'Não informado',
                '{{prescricao_resumo}}': case.prescription_summary[:100] + '...' if case.prescription_summary and len(case.prescription_summary) > 100 else case.prescription_summary or 'Não informado',
            },
        }
    
    def _replace_placeholders_in_document(self, document, case, benefits):
        """
        Substitui placeholders {{variavel}} no documento pelos dados reais
        Procura em todo o documento: parágrafos, tabelas, cabeçalhos e rodapés
        """
        replacements = {
            # Dados do Cliente
            '{{cliente_nome}}': case.client.name if case.client else '',
            '{{cliente_cnpj}}': case.client.cnpj if case.client else '',
            '{{cliente_endereco}}': self._format_address(case.client) if case.client else '',
            '{{cliente_cidade}}': case.client.city if case.client else '',
            '{{cliente_estado}}': case.client.state if case.client else '',
            
            # Dados do Caso
            '{{caso_titulo}}': case.title,
            '{{caso_tipo}}': case.case_type,
            '{{caso_numero}}': str(case.id),
            
            # Dados FAP
            '{{fap_motivo}}': self._get_fap_reason_text(case.fap_reason) if case.fap_reason else '',
            '{{ano_inicial_fap}}': str(case.fap_start_year) if case.fap_start_year else '',
            '{{ano_final_fap}}': str(case.fap_end_year) if case.fap_end_year else '',
            '{{anos_fap}}': self._format_years_range(case.fap_start_year, case.fap_end_year),
            
            # Dados de Benefícios
            '{{total_beneficios}}': str(len(benefits)),
            '{{lista_beneficios}}': self._format_benefits_list(benefits),
            '{{quantidade_acidentes}}': str(len(benefits)),
            '{{quantidade_acidentes_extenso}}': self._number_to_words(len(benefits)),
            '{{titulo_acidente_trajeto}}': 'Acidente de Trajeto' if len(benefits) == 1 else 'Acidentes de Trajeto',
            '{{nome_segurado_exemplo}}': benefits[0].insured_name if benefits else '',
            '{{nit_segurado_exemplo}}': benefits[0].insured_nit if benefits else '',
            '{{data_acidente_segurado_exemplo}}': benefits[0].accident_date.strftime('%d/%m/%Y') if benefits and benefits[0].accident_date else '',
            '{{data_inicio_beneficio_segurado_exemplo}}': benefits[0].data_inicio_beneficio.strftime('%d/%m/%Y') if benefits and benefits[0].data_inicio_beneficio else '',
            '{{data_fim_beneficio_segurado_exemplo}}': benefits[0].data_fim_beneficio.strftime('%d/%m/%Y') if benefits and benefits[0].data_fim_beneficio else '',
            '{{numero_beneficio}}': benefits[0].benefit_number if benefits else '',
            '{{lesao_segurado_exemplo}}': '',  # Campo adicional
            '{{resumo_evento_segurado_exemplo}}': '',  # Campo adicional
            '{{data_afastamento_trabalho}}': '',  # Campo adicional
            '{{data_laudo_medico_segurado_exemplo}}': '',  # Campo adicional
            '{{numero_cat}}': benefits[0].numero_cat if benefits and benefits[0].numero_cat else '',
            '{{numero_bo}}': benefits[0].numero_bo if benefits and benefits[0].numero_bo else '',
            
            # Valores
            '{{valor_causa}}': self._format_currency(case.value_cause) if case.value_cause else '',
            '{{valor_causa_extenso}}': self._value_to_words(case.value_cause) if case.value_cause else '',
            
            # Datas
            '{{data_ajuizamento}}': case.filing_date.strftime('%d/%m/%Y') if case.filing_date else '',
            '{{data_atual}}': datetime.now().strftime('%d/%m/%Y'),
            '{{mes_ano_atual}}': datetime.now().strftime('%B de %Y'),
            '{{vigencia_fap}}': self._format_years_range(case.fap_start_year, case.fap_end_year),
            
            # Imagens (placeholders - inserção real requer lógica adicional)
            '{{Imagem_cat}}': '[IMAGEM_CAT]',
            '{{imagem_fap}}': '[IMAGEM_FAP]',
            '{{imagem_info_beneficiario}}': '[IMAGEM_INFO_BENEFICIARIO]',
            '{{imagem_declaracao_beneficio}}': '[IMAGEM_DECLARACAO_BENEFICIO]',
            '{{imagem_inss_beneficiario}}': '[IMAGEM_INSS_BENEFICIARIO]',
            '{{imagem_vigencia_beneficio}}': '[IMAGEM_VIGENCIA_BENEFICIO]',
            
            # Dados da Vara
            '{{vara_nome}}': case.court.vara_name if case.court else '',
            '{{vara_cidade}}': case.court.city if case.court else '',
            '{{vara_estado}}': case.court.state if case.court else '',
            '{{vara_completo}}': self._format_court(case.court) if case.court else '',
            
            # Resumos
            '{{fatos_resumo}}': case.facts_summary or '',
            '{{teses_resumo}}': case.thesis_summary or '',
            '{{prescricao_resumo}}': case.prescription_summary or '',
        }
        
        # Substituir em parágrafos do corpo do documento
        for paragraph in document.paragraphs:
            self._replace_in_paragraph(paragraph, replacements)
        
        # Substituir em tabelas
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, replacements)
        
        # Substituir em cabeçalhos
        for section in document.sections:
            header = section.header
            for paragraph in header.paragraphs:
                self._replace_in_paragraph(paragraph, replacements)
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_in_paragraph(paragraph, replacements)
        
        # Substituir em rodapés
        for section in document.sections:
            footer = section.footer
            for paragraph in footer.paragraphs:
                self._replace_in_paragraph(paragraph, replacements)
            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_in_paragraph(paragraph, replacements)
    
    def _replace_in_paragraph(self, paragraph, replacements):
        """
        Substitui placeholders em um parágrafo específico
        Lida com placeholders que podem estar divididos entre múltiplos runs
        """
        # Primeiro, verificar se há placeholders no texto completo
        full_text = paragraph.text
        
        # Para cada placeholder
        for key, value in replacements.items():
            if key in full_text:
                # Abordagem 1: Tentar substituir run por run (mais seguro para formatação)
                replaced = False
                for run in paragraph.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, value)
                        replaced = True
                
                # Abordagem 2: Se o placeholder está dividido entre runs
                if not replaced and key in full_text:
                    # Reconstruir o parágrafo mantendo apenas o primeiro run
                    if paragraph.runs:
                        # Pegar formatação do primeiro run
                        first_run = paragraph.runs[0]
                        
                        # Substituir texto no parágrafo completo
                        new_text = full_text.replace(key, value)
                        
                        # Limpar todos os runs
                        for run in paragraph.runs:
                            run.text = ''
                        
                        # Adicionar texto no primeiro run
                        first_run.text = new_text
                        full_text = new_text
    
    def _format_address(self, client):
        """Formata endereço completo do cliente"""
        if not client:
            return ''
        parts = []
        if client.street:
            parts.append(client.street)
        if client.number:
            parts.append(f"nº {client.number}")
        if client.district:
            parts.append(client.district)
        if client.city and client.state:
            parts.append(f"{client.city}/{client.state}")
        if client.zip_code:
            parts.append(f"CEP {client.zip_code}")
        return ', '.join(parts)
    
    def _format_currency(self, value):
        """Formata valor monetário"""
        return f"R$ {value:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
    
    def _format_years_range(self, start_year, end_year):
        """Formata range de anos para FAP"""
        if start_year and end_year:
            if start_year == end_year:
                return str(start_year)
            return f"{start_year} a {end_year}"
        return ''
    
    def _format_benefits_list(self, benefits):
        """Formata lista de benefícios para inserção no texto"""
        if not benefits:
            return ''
        lines = []
        for idx, benefit in enumerate(benefits, start=1):
            line = f"{idx}. NB {benefit.benefit_number} - {benefit.insured_name}"
            if benefit.benefit_type:
                line += f" ({benefit.benefit_type})"
            lines.append(line)
        return '\n'.join(lines)
    
    def _format_court(self, court):
        """Formata nome completo da vara"""
        if not court:
            return ''
        parts = []
        if court.vara_name:
            parts.append(court.vara_name)
        if court.city and court.state:
            parts.append(f"{court.city}/{court.state}")
        return ' - '.join(parts)
    
    def _value_to_words(self, value):
        """Converte valor numérico para extenso (simplificado)"""
        # Implementação simplificada - pode ser expandida
        return f"{value:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
    
    def _number_to_words(self, number):
        """Converte número inteiro para extenso (português)"""
        numbers = {
            0: 'zero', 1: 'um', 2: 'dois', 3: 'três', 4: 'quatro', 5: 'cinco',
            6: 'seis', 7: 'sete', 8: 'oito', 9: 'nove', 10: 'dez',
            11: 'onze', 12: 'doze', 13: 'treze', 14: 'quatorze', 15: 'quinze',
            16: 'dezesseis', 17: 'dezessete', 18: 'dezoito', 19: 'dezenove', 20: 'vinte',
            30: 'trinta', 40: 'quarenta', 50: 'cinquenta', 60: 'sessenta',
            70: 'setenta', 80: 'oitenta', 90: 'noventa', 100: 'cem'
        }
        
        if number in numbers:
            return numbers[number]
        elif number < 100:
            tens = (number // 10) * 10
            units = number % 10
            return f"{numbers[tens]} e {numbers[units]}"
        else:
            return str(number)  # Retorna número se > 100 (pode expandir se necessário)
    
    def find_placeholders_in_document(self, document):
        """
        Método de debug: encontra todos os placeholders no documento
        
        Args:
            document: Documento DOCX carregado
            
        Returns:
            list: Lista de placeholders únicos encontrados
        """
        import re
        placeholders = set()
        pattern = r'\{\{[^}]+\}\}'
        
        # Procurar em parágrafos
        for paragraph in document.paragraphs:
            matches = re.findall(pattern, paragraph.text)
            placeholders.update(matches)
        
        # Procurar em tabelas
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        matches = re.findall(pattern, paragraph.text)
                        placeholders.update(matches)
        
        # Procurar em cabeçalhos e rodapés
        for section in document.sections:
            for paragraph in section.header.paragraphs:
                matches = re.findall(pattern, paragraph.text)
                placeholders.update(matches)
            for paragraph in section.footer.paragraphs:
                matches = re.findall(pattern, paragraph.text)
                placeholders.update(matches)
        
        return sorted(list(placeholders))
    
    def _get_fap_reason_text(self, fap_reason_code):
        """Converte código do motivo FAP em texto legível"""
        reasons = {
            'inclusao_indevida_trajeto': 'Inclusão indevida de benefício de trajeto',
            'erro_material_cat': 'Erro material no preenchimento da CAT',
            'cat_trajeto_extemporanea': 'CAT de trajeto transmitida extemporaneamente'
        }
        return reasons.get(fap_reason_code, fap_reason_code)
    
    def _add_benefits_to_tables(self, document, case, benefits):
        """
        Adiciona linhas com dados dos benefícios nas tabelas do documento
        Procura por tabelas que contenham marcadores específicos ou cabeçalhos
        
        Args:
            document: Documento DOCX
            case: Objeto Case com dados do caso (incluindo anos FAP)
            benefits: Lista de objetos CaseBenefit
        """
        if not benefits:
            return
        
        # Procurar a tabela de benefícios (identificada por cabeçalhos específicos)
        benefit_table = None
        for table in document.tables:
            # Verificar se é a tabela de benefícios pelos cabeçalhos
            if len(table.rows) > 0:
                first_row_text = ' '.join([cell.text.lower() for cell in table.rows[0].cells])
                # Procurar por palavras-chave que indicam tabela de benefícios
                keywords = ['benefício', 'nb', 'segurado', 'nit', 'acidente']
                if any(keyword in first_row_text for keyword in keywords):
                    benefit_table = table
                    break
        
        # Se não encontrou tabela específica, usar a primeira tabela com 6+ colunas
        if benefit_table is None:
            for table in document.tables:
                if len(table.rows) > 0 and len(table.rows[0].cells) >= 6:
                    benefit_table = table
                    break
        
        # Adicionar benefícios na tabela encontrada
        if benefit_table:
            for idx, benefit in enumerate(benefits, start=1):
                new_row = benefit_table.add_row()
                cells = new_row.cells
                
                # Preencher células com dados do benefício
                # Estrutura: Item | Vigências do FAP | CNPJ | Empregado(a) | NIT | Tipo | Benefício | CAT
                if len(cells) >= 1:
                    # Item
                    self._format_table_cell(cells[0], str(idx))
                if len(cells) >= 2:
                    # Vigências do FAP (anos inicial e final combinados)
                    vigencia = self._format_years_range(case.fap_start_year, case.fap_end_year)
                    self._format_table_cell(cells[1], vigencia)
                if len(cells) >= 3:
                    # CNPJ do cliente
                    cnpj = case.client.cnpj if case.client else ''
                    self._format_table_cell(cells[2], cnpj)
                if len(cells) >= 4:
                    # Empregado(a) - Nome do Segurado
                    self._format_table_cell(cells[3], benefit.insured_name or '')
                if len(cells) >= 5:
                    # NIT
                    self._format_table_cell(cells[4], benefit.insured_nit or '')
                if len(cells) >= 6:
                    # Tipo de Benefício
                    self._format_table_cell(cells[5], benefit.benefit_type or '')
                if len(cells) >= 7:
                    # Benefício - Número do Benefício
                    self._format_table_cell(cells[6], benefit.benefit_number or '')
                if len(cells) >= 8:
                    # CAT - Data do Acidente
                    cat_data = benefit.accident_date.strftime('%d/%m/%Y') if benefit.accident_date else ''
                    self._format_table_cell(cells[7], cat_data)
    
    def _format_table_cell(self, cell, text):
        """
        Formata uma célula de tabela com estilo específico:
        - Fonte: Avenir Next LT Pro
        - Tamanho: 7pt
        - Alinhamento: Centralizado
        
        Args:
            cell: Objeto Cell do python-docx
            text: Texto a ser inserido na célula
        """
        # Limpar conteúdo existente
        cell.text = ''
        
        # Adicionar parágrafo com formatação
        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Adicionar run com texto formatado
        run = paragraph.add_run(text)
        run.font.name = 'Avenir Next LT Pro'
        run.font.size = Pt(7)


# =====================================
# Script de Teste e Desenvolvimento
# =====================================
if __name__ == "__main__":
    """Script de teste - roda apenas quando executado diretamente"""
    print("=== Teste do AgentDocumentGenerator ===\n")
    
    # Carregar template de exemplo
    document = Document("modelo_acidente_trajeto.docx")
    document.add_heading("Modelo de Documento - TESTE", 0)
    
    # Testar iteração sobre tabelas
    print("Conteúdo das tabelas:")
    for table_idx, table in enumerate(document.tables):
        print(f"\nTabela {table_idx + 1}:")
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        print(f"  {paragraph.text}")
        
        # Adicionar linha de teste
        new_row = table.add_row()
        mock_data = {
            'id': 1,
            'benefit_number': '123456789',
            'insured_name': 'João da Silva Santos',
            'insured_nit': '123.45678.90-1',
            'accident_date': '15/03/2024',
            'benefit_type': 'B91'
        }
        
        for i, cell in enumerate(new_row.cells):
            if i == 0: 
                cell.text = str(mock_data['id'])
            elif i == 1: 
                cell.text = mock_data['benefit_number']
            elif i == 2: 
                cell.text = mock_data['insured_name']
            elif i == 3: 
                cell.text = mock_data['insured_nit']
            elif i == 4: 
                cell.text = mock_data['accident_date']
            elif i == 5: 
                cell.text = mock_data['benefit_type']
            else: 
                cell.text = f"Dado {i+1}"
        print(f"  ✓ Linha de teste adicionada")
    
    # Buscar placeholders
    print("\nPlaceholders encontrados:")
    for p in document.paragraphs:
        if "{{" in p.text and "}}" in p.text:
            print(f"  {p.text}")
    
    # Salvar
    output_path = "modelo_acidente_trajeto_edit.docx"
    document.save(output_path)
    print(f"\n✓ Documento salvo em: {output_path}")
    
    # Instruções
    print("\n=== Uso em Produção ===")
    print("""
Para usar a classe em produção:

from agent_document_generator import AgentDocumentGenerator

agent = AgentDocumentGenerator()

# Forma 1: Seleção automática de template baseada no fap_reason
document = agent.generate_fap_petition(case_id=123)

# Forma 2: Template específico (override)
document = agent.generate_fap_petition(
    case_id=123,
    template_path="templates_docx/modelo_acidente_trajeto.docx"
)

document.save("peticao_gerada.docx")

Templates disponíveis:
- inclusao_indevida_trajeto → modelo_acidente_trajeto.docx
- erro_material_cat → modelo_acidente_trajeto_erro_material.docx
- cat_trajeto_extemporanea → modelo_acidente_trajeto_extemporanea.docx
""")
