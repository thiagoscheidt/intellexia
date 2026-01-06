from io import BytesIO
from openai import OpenAI
from app.prompts.document_reader_prompt import DocumentReaderPrompt
from langchain_openai import ChatOpenAI
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches
from app.models import Case, CaseBenefit

_ = load_dotenv()


class AgentDocumentGenerator:
    """
    Agente especializado para geração de petições FAP usando modelos DOCX
    Processa templates Word e preenche com dados do banco de dados
    """
    model = None
    prompt = None

    def __init__(self, model_name="gpt-4o"):
        self.model = ChatOpenAI(model=model_name)
        self.prompt = DocumentReaderPrompt()

    def generate(self, file_id):
        """Gera resumo de documento usando IA"""
        result = self.model.invoke(
            [
                {"role": "system", "content": self.prompt.prompt_template()},
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": "Analise o documento e gere um resumo objetivo.",
                        },
                        {"type": "file", "file_id": file_id},
                    ],
                },
            ]
        )
        return result.content
    
    def generate_fap_petition(self, case_id, template_path="modelo_acidente_trajeto.docx"):
        """
        Gera petição FAP preenchendo template DOCX com dados do banco
        
        Args:
            case_id: ID do caso no banco de dados
            template_path: Caminho para o template DOCX
            
        Returns:
            Document: Documento DOCX preenchido
        """
        # Buscar dados do caso
        case = Case.query.get(case_id)
        if not case:
            raise ValueError(f"Caso {case_id} não encontrado")
        
        # Buscar benefícios do caso
        benefits = CaseBenefit.query.filter_by(case_id=case_id).all()
        
        # Carregar template
        document = Document(template_path)
        
        # Preencher campos do template com dados do caso
        self._replace_placeholders_in_document(document, case, benefits)
        
        # Adicionar benefícios nas tabelas
        self._add_benefits_to_tables(document, benefits)
        
        return document
    
    def _replace_placeholders_in_document(self, document, case, benefits):
        """Substitui placeholders {{variavel}} no documento pelos dados reais"""
        replacements = {
            '{{cliente_nome}}': case.client.name if case.client else '',
            '{{cliente_cnpj}}': case.client.cnpj if case.client else '',
            '{{caso_titulo}}': case.title,
            '{{caso_tipo}}': case.case_type,
            '{{fap_motivo}}': self._get_fap_reason_text(case.fap_reason) if case.fap_reason else '',
            '{{ano_inicial_fap}}': str(case.fap_start_year) if case.fap_start_year else '',
            '{{ano_final_fap}}': str(case.fap_end_year) if case.fap_end_year else '',
            '{{total_beneficios}}': str(len(benefits)),
            '{{valor_causa}}': f"R$ {case.value_cause:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.') if case.value_cause else '',
            '{{data_ajuizamento}}': case.filing_date.strftime('%d/%m/%Y') if case.filing_date else '',
            '{{vara_nome}}': case.court.vara_name if case.court else '',
            '{{vara_cidade}}': case.court.city if case.court else '',
            '{{vara_estado}}': case.court.state if case.court else '',
        }
        
        # Substituir em parágrafos
        for paragraph in document.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    # Preservar formatação
                    for run in paragraph.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, value)
        
        # Substituir em tabelas
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in paragraph.text:
                                for run in paragraph.runs:
                                    if key in run.text:
                                        run.text = run.text.replace(key, value)
    
    def _get_fap_reason_text(self, fap_reason_code):
        """Converte código do motivo FAP em texto legível"""
        reasons = {
            'inclusao_indevida_trajeto': 'Inclusão indevida de benefício de trajeto',
            'erro_material_cat': 'Erro material no preenchimento da CAT',
            'cat_trajeto_extemporanea': 'CAT de trajeto transmitida extemporaneamente'
        }
        return reasons.get(fap_reason_code, fap_reason_code)
    
    def _add_benefits_to_tables(self, document, benefits):
        """Adiciona linhas com dados dos benefícios nas tabelas do documento"""
        if not benefits:
            return
        
        # Procurar tabelas que possam conter benefícios
        # (normalmente a primeira tabela após certos marcadores)
        for table in document.tables:
            # Adicionar uma linha para cada benefício
            for idx, benefit in enumerate(benefits, start=1):
                new_row = table.add_row()
                
                # Preencher células com dados do benefício
                if len(new_row.cells) >= 6:
                    new_row.cells[0].text = str(idx)  # ID/Número sequencial
                    new_row.cells[1].text = benefit.benefit_number or ''
                    new_row.cells[2].text = benefit.insured_name or ''
                    new_row.cells[3].text = benefit.insured_nit or ''
                    new_row.cells[4].text = benefit.accident_date.strftime('%d/%m/%Y') if benefit.accident_date else ''
                    new_row.cells[5].text = benefit.benefit_type or ''
                    
                    # Se houver mais colunas, preencher com dados adicionais
                    if len(new_row.cells) > 6:
                        new_row.cells[6].text = benefit.error_reason or ''
                    if len(new_row.cells) > 7:
                        new_row.cells[7].text = benefit.accident_company_name or ''


# =====================================
# Script de Teste e Desenvolvimento
# =====================================
if __name__ == "__main__":
    """
    Script de teste para desenvolvimento do AgentDocumentGenerator
    """
    
    # Carregar template de exemplo
    document = Document("modelo_acidente_trajeto.docx")
document.add_heading("Modelo de Documento", 0)


# for table in document.tables:
#     for row in table.rows:
#         for cell in row.cells:
#             for paragraph in cell.paragraphs:
#                 print(paragraph.text)

for table_idx, table in enumerate(document.tables):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                print(paragraph.text)
    
    # Adiciona nova linha na tabela com dados do sistema
    new_row = table.add_row()
    
    # Simular dados de um caso real do sistema
    # Em produção, esses dados viriam do banco de dados
    mock_case_data = {
        'id': 1,
        'benefit_number': '123456789',
        'insured_name': 'João da Silva Santos',
        'insured_nit': '123.45678.90-1',
        'accident_date': '15/03/2024',
        'benefit_type': 'B91',
        'status': 'Ativo',
        'value': 'R$ 1.500,00',
        'company': 'Empresa XYZ Ltda',
        'fap_reason': 'Trajeto de Ida ao Trabalho'
    }
    
    for i, cell in enumerate(new_row.cells):
        if i == 0:
            # Primeira coluna: ID
            cell.text = str(mock_case_data['id'])
        elif i == 1:
            cell.text = mock_case_data['benefit_number']
        elif i == 2:
            cell.text = mock_case_data['insured_name']
        elif i == 3:
            cell.text = mock_case_data['insured_nit']
        elif i == 4:
            cell.text = mock_case_data['accident_date']
        elif i == 5:
            cell.text = mock_case_data['benefit_type']
        else:
            # Colunas adicionais
            cell.text = mock_case_data.get('status', f"Dado coluna {i+1}")

for p in document.paragraphs:
    if p.text.startswith("{{") and p.text.endswith("}}"):
        print("\n", p.text)
    # for run in p.runs:
    #     print(run.text)
# agent = AgentDocumentGenerator()
# response = agent.generate()

document.save("modelo_acidente_trajeto_edit.docx")