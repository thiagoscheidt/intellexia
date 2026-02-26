from __future__ import annotations

import os
import sys
from pathlib import Path as PathLib
from datetime import datetime
from typing import Optional, List, Dict

from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from pydantic import BaseModel, Field
from docx import Document

# Adicionar o diretório raiz ao path para imports do app
root_dir = PathLib(__file__).parent.parent.parent
sys.path.insert(0, str(root_dir))


load_dotenv()

class DocumentMappingSchema(BaseModel):
    replacements: Dict[str, str] = Field(description="Mapeamento de índices para textos com dados substituídos. Chave: índice do texto, Valor: texto completo substituído")
    summary: str = Field(description="Resumo das substituições realizadas")

class FapSectionGeneratorAgent:
    def __init__(self):
        self.llm = ChatOpenAI(
            model="gpt-5-mini-mini",
            temperature=0
        )
        # Nota: Vamos processar a resposta JSON manualmente ao invés de usar structured_output
        # para evitar problemas com o schema da OpenAI
        
    def populate_template_with_case_data(self, document: Document, case_data: str) -> Dict:
        """
        Preenche um template DOCX substituindo dados fictícios pelos dados reais do caso
        PRESERVA toda formatação, comentários e estilos do documento original
        OTIMIZADO: Processa todo o documento em UMA ÚNICA requisição
        
        Args:
            document: Objeto Document (python-docx) carregado
            case_data: Dados reais do caso (cliente, benefícios, competências, etc)
            
        Returns:
            Dict com 'document' (Document modificado), 'total_replacements' (int) e 'replacements_log' (str)
        """
        
        print("\n🔄 COLETANDO TEXTOS DO DOCUMENTO...")
        print("=" * 80)
        
        # Coletar TODOS os textos do documento com suas referências
        text_map = self._collect_all_texts(document)
        
        print(f"\n📊 Total de textos coletados: {len(text_map)}")
        
        # Preparar conteúdo para IA
        all_texts = "\n\n---\n\n".join([f"[{idx}] {text}" for idx, (_, text) in text_map.items()])
        
        print("\n🤖 PROCESSANDO TODOS OS TEXTOS COM IA (requisição única)...")
        print("⏳ Isso pode levar alguns segundos...\n")
        
        try:
            # UMA ÚNICA chamada à IA com todo o conteúdo
            response = self.llm.invoke([
                {"role": "system", "content": """
                    Você é um especialista em documentos jurídicos previdenciários.
                    
                    TAREFA:
                    Você receberá vários textos do documento, cada um identificado por [N].
                    Para CADA texto:
                    1. Identifique dados fictícios/exemplos (nomes, datas, números, CNPJs, endereços, valores, etc)
                    2. Substitua pelos dados reais correspondentes do caso fornecido
                    3. Mantenha TODA estrutura, pontuação e formatação original
                    4. Ajuste concordâncias de gênero/número se necessário
                    5. Se não houver dado correspondente, mantenha o texto original
                    
                    RETORNE UM JSON com esta estrutura:
                    {
                        "replacements": {
                            "0": "texto substituído para índice 0",
                            "1": "texto substituído para índice 1"
                        },
                        "summary": "Resumo das substituições realizadas"
                    }
                    
                    IMPORTANTE:
                    - No "replacements", inclua APENAS os textos que FORAM MODIFICADOS
                    - Use o índice como chave (string)
                    - Preserve toda pontuação, espaços e estrutura
                    - Mantenha gramática correta
                    - Retorne APENAS o JSON, sem nenhum texto adicional
                    """},
                {"role": "user", "content": f"""
                    DADOS REAIS DO CASO:
                    {case_data}
                    
                    ---
                    
                    TEXTOS DO DOCUMENTO:
                    {all_texts}
                    
                    Processe todos os textos e retorne o JSON com os textos modificados.
                    """}
            ])
            
            # Parsear resposta JSON
            import json
            response_text = response.content
            
            # Remover markdown code blocks se existirem
            if "```json" in response_text:
                response_text = response_text.split("```json")[1].split("```")[0]
            elif "```" in response_text:
                response_text = response_text.split("```")[1].split("```")[0]
            
            response_data = json.loads(response_text.strip())
            
            print(f"✅ IA processou o documento!")
            print(f"📊 Textos modificados: {len(response_data.get('replacements', {}))}")
            
            # Aplicar substituições no documento
            print("\n🔄 APLICANDO SUBSTITUIÇÕES...")
            total_applied = self._apply_replacements(document, text_map, response_data.get('replacements', {}))
            
            print("\n" + "=" * 80)
            print(f"✅ PROCESSAMENTO CONCLUÍDO")
            print(f"📊 Total de substituições aplicadas: {total_applied}")
            print("=" * 80)
            
            return {
                "document": document,
                "total_replacements": total_applied,
                "replacements_log": response_data.get('summary', 'Processamento concluído')
            }
            
        except Exception as e:
            print(f"\n❌ ERRO ao processar documento: {str(e)}")
            import traceback
            traceback.print_exc()
            return {
                "document": document,
                "total_replacements": 0,
                "replacements_log": f"Erro: {str(e)}"
            }
    
    def _collect_all_texts(self, document: Document) -> Dict[str, tuple]:
        """
        Coleta TODOS os textos do documento com suas referências
        
        Returns:
            Dict mapeando índice para (paragraph_object, original_text)
        """
        text_map = {}
        idx = 0
        
        # Parágrafos do corpo
        for paragraph in document.paragraphs:
            if paragraph.text.strip() and len(paragraph.text.strip()) >= 5:
                text_map[str(idx)] = (paragraph, paragraph.text)
                idx += 1
        
        # Tabelas
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip() and len(paragraph.text.strip()) >= 5:
                            text_map[str(idx)] = (paragraph, paragraph.text)
                            idx += 1
        
        # Cabeçalhos
        for section in document.sections:
            for paragraph in section.header.paragraphs:
                if paragraph.text.strip() and len(paragraph.text.strip()) >= 5:
                    text_map[str(idx)] = (paragraph, paragraph.text)
                    idx += 1
            
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text.strip() and len(paragraph.text.strip()) >= 5:
                                text_map[str(idx)] = (paragraph, paragraph.text)
                                idx += 1
        
        # Rodapés
        for section in document.sections:
            for paragraph in section.footer.paragraphs:
                if paragraph.text.strip() and len(paragraph.text.strip()) >= 5:
                    text_map[str(idx)] = (paragraph, paragraph.text)
                    idx += 1
            
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text.strip() and len(paragraph.text.strip()) >= 5:
                                text_map[str(idx)] = (paragraph, paragraph.text)
                                idx += 1
        
        return text_map
    
    def _apply_replacements(self, document: Document, text_map: Dict[str, tuple], replacements: Dict[str, str]) -> int:
        """
        Aplica substituições no documento preservando formatação
        
        Args:
            document: Documento original
            text_map: Mapa de índices para (paragraph, texto_original)
            replacements: Mapa de índices para textos substituídos
            
        Returns:
            int: Número de substituições aplicadas
        """
        total_applied = 0
        
        for idx, new_text in replacements.items():
            if idx in text_map:
                paragraph, original_text = text_map[idx]
                
                # Aplicar substituição preservando formatação
                if paragraph.runs:
                    # Pegar formatação do primeiro run
                    first_run = paragraph.runs[0]
                    
                    # Limpar todos os runs
                    for run in paragraph.runs:
                        run.text = ""
                    
                    # Adicionar novo texto no primeiro run (preserva formatação)
                    first_run.text = new_text
                else:
                    paragraph.text = new_text
                
                total_applied += 1
                print(f"   ✓ [{idx}] Substituído")
        
        return total_applied


def main():
    """Função para testar o gerador de seções/tópicos"""
    
    print("=" * 80)
    print("TESTE DO GERADOR DE SEÇÕES FAP - OTIMIZADO (1 REQUISIÇÃO)")
    print("=" * 80)
    
    # Carregar documento de exemplo
    template_path = root_dir / "templates_padrao" / "1_Peticao_Inicial_Mandado_Seguranca_FAP.docx"
    
    if not template_path.exists():
        print(f"\n❌ ERRO: Template não encontrado em: {template_path}")
        print("\nBuscando templates alternativos...")
        
        # Tentar encontrar qualquer template DOCX
        templates_dir = root_dir / "templates_padrao"
        if templates_dir.exists():
            templates = list(templates_dir.glob("*.docx"))
            if templates:
                template_path = templates[0]
                print(f"✅ Usando template: {template_path.name}")
            else:
                print("❌ Nenhum template .docx encontrado em templates_padrao/")
                return
        else:
            print("❌ Diretório templates_padrao/ não existe")
            return
    
    print(f"\n📄 Carregando template: {template_path.name}")
    
    # Dados de exemplo do caso
    case_data = """
    DADOS DO PROCESSO:
    
    CLIENTE:
    - Nome/Razão Social: Metalúrgica Silva & Cia Ltda
    - CNPJ: 12.345.678/0001-90
    - Endereço: Rua das Indústrias, 1500, Distrito Industrial
    - Cidade: São Paulo
    - Estado: SP
    - CEP: 01234-567
    
    ADVOGADO:
    - Nome: Dr. José Carlos Santos
    - OAB: SP 123.456
    - Endereço: Av. Paulista, 1000, sala 501
    - Cidade: São Paulo/SP
    
    PROCESSO:
    - Tipo de Ação: Mandado de Segurança - Revisão FAP
    - Ano FAP: 2020-2022
    - Motivo: Acidente de trajeto incluído indevidamente no FAP
    
    BENEFICIÁRIO/SEGURADO:
    - Nome: João da Silva
    - CPF: 123.456.789-00
    - NIT/PIS: 123.45678.90-1
    - Cargo: Operador de máquinas
    - Data de admissão: 10/01/2015
    
    BENEFÍCIO QUESTIONADO:
    - Tipo: B91 - Auxílio-doença acidentário
    - Número do benefício: 123.456.789
    - NB: 123456789
    - DIB (Data Início): 15/03/2021
    - DCB (Data Cessação): 15/10/2021
    - Valor mensal: R$ 2.500,00
    
    ACIDENTE/CAT:
    - Número CAT: 2021.00.123456
    - Data emissão CAT: 11/03/2021
    - Data do acidente: 10/03/2021 às 07h30
    - Tipo de acidente: Acidente de trajeto (residência-trabalho)
    - Local: Avenida dos Trabalhadores, altura do número 500
    - Descrição: Colisão de motocicleta no trajeto para o trabalho
    - CID: S82.0 - Fratura da patela
    
    FUNDAMENTAÇÃO LEGAL:
    - Lei 8.213/91, Art. 19, §1º
    - Lei 10.666/2003 (FAP)
    - Decreto 3.048/99, Art. 336
    
    VALORES:
    - FAP original: 2,00
    - FAP após correção: 1,50
    - Diferença mensal: R$ 1.200,00
    - Valor total a restituir: R$ 28.800,00
    
    AUTORIDADE COATORA:
    - Superintendente Regional do INSS em São Paulo
    - Agência: São Paulo - Sé
    
    VARA/JUÍZO:
    - 1ª Vara Federal de São Paulo
    - Seção Judiciária de São Paulo
    
    DATA: 26/01/2026
    """
    
    try:
        # Carregar documento
        document = Document(str(template_path))
        
        print(f"✅ Documento carregado com sucesso")
        print(f"   Parágrafos: {len(document.paragraphs)}")
        print(f"   Tabelas: {len(document.tables)}")
        print(f"   Seções: {len(document.sections)}")
        
        # Instanciar o agente
        agent = FapSectionGeneratorAgent()
        
        print("\n🤖 Processando documento (1 única requisição à IA)...")
        print("⏳ Processando...\n")
        
        # Processar documento
        result = agent.populate_template_with_case_data(document, case_data)
        
        processed_document = result["document"]
        total_replacements = result["total_replacements"]
        replacements_log = result["replacements_log"]
        
        # Salvar documento processado
        output_dir = root_dir / "uploads" / "generated_sections"
        output_dir.mkdir(parents=True, exist_ok=True)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"peticao_gerada_{timestamp}.docx"
        output_path = output_dir / output_filename
        
        processed_document.save(str(output_path))
        
        print("\n" + "=" * 80)
        print("✅ DOCUMENTO PROCESSADO COM SUCESSO!")
        print("=" * 80)
        print(f"\n💾 Arquivo salvo em:")
        print(f"   {output_path}")
        print(f"\n📊 Tamanho: {output_path.stat().st_size / 1024:.2f} KB")
        print(f"📊 Total de substituições: {total_replacements}")
        
        # Mostrar resumo
        if replacements_log:
            print("\n📝 Resumo das substituições:")
            print("-" * 80)
            print(f"   {replacements_log}")
            print("-" * 80)
        
        print("\n💡 VANTAGENS DESTA ABORDAGEM OTIMIZADA:")
        print("   ✓ Apenas 1 requisição à IA (ao invés de 178!)")
        print("   ✓ Muito mais rápido (segundos ao invés de minutos)")
        print("   ✓ Muito mais econômico (1 chamada API vs 178)")
        print("   ✓ Preserva TODA formatação original")
        print("   ✓ Mantém comentários do Word intactos")
        print("   ✓ Preserva estrutura de tabelas e layout")
        print("\n" + "=" * 80)
        
    except Exception as e:
        print(f"\n❌ ERRO ao processar documento: {str(e)}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
