"""Agente para padronizar e exportar documentos para DOCX no padrГЈo do escritГіrio.

Responsabilidades:
1. (Opcional) Normalizar o texto do documento com LLM mini para reduzir ruГ­do de formataГ§ГЈo.
2. Aplicar template templates_padrao/modelo_documento.docx (cabeГ§alho/rodapГ©/margens).
3. Gerar stream DOCX pronto para download.
"""

from __future__ import annotations

import os
import re
import shutil
import tempfile
import time
from io import BytesIO
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from langchain_openai import ChatOpenAI
from pydantic import BaseModel, Field

from app.agents.config import DEFAULT_MODEL_MINI
from app.services.token_usage_service import TokenUsageService


class NormalizedDocxText(BaseModel):
    """SaГ­da de normalizaГ§ГЈo para exportaГ§ГЈo em DOCX."""

    normalized_text: str = Field(
        description=(
            "Texto final normalizado para DOCX, preservando integralmente o conteГєdo jurГ­dico. "
            "Pode usar marcaГ§Гµes simples em Markdown (#, ##, ###, ---)."
        )
    )


_SYSTEM_PROMPT = (
    "VocГЄ normaliza textos jurГ­dicos para exportaГ§ГЈo DOCX. "
    "Preserve conteГєdo e sentido jurГ­dico integralmente, sem resumir nem omitir seГ§Гµes. "
    "Apenas padronize quebras de linha, tГ­tulos e legibilidade. "
    "Nunca altere pedidos, fundamentos, valores, NBs, NITs, CNPJs ou nГєmeros processuais. "
    "Preserve tabelas markdown com pipes exatamente como tabelas estruturadas. "
    "Linhas no formato [Anexo: ...] sГЈo referГЄncias de prova вЂ” preserve-as exatamente, em linha prГіpria."
)

_BOLD_MARKDOWN_RE = re.compile(r"(\*\*([^*]+)\*\*|__([^_]+)__)")
_PLAIN_HEADING_RE = re.compile(
    r"^(?:"
    r"(?:[IVXLCDM]{1,8}|\d+(?:\.\d+)*)\.\s+[A-ZГЃГ‰ГЌГ“ГљГЂГ‚ГЉГ”ГѓГ•Г‡0-9][A-ZГЃГ‰ГЌГ“ГљГЂГ‚ГЉГ”ГѓГ•Г‡0-9\s\-/,&().]{2,}"
    r"|"
    r"[A-ZГЃГ‰ГЌГ“ГљГЂГ‚ГЉГ”ГѓГ•Г‡0-9][A-ZГЃГ‰ГЌГ“ГљГЂГ‚ГЉГ”ГѓГ•Г‡0-9\s\-/,&().]{3,}"
    r")$"
)


class OfficeDocxExportAgent:
    """Padroniza texto e exporta DOCX com template do escritГіrio."""

    def __init__(self, model_name: Optional[str] = None, temperature: float = 0.0):
        self.model_name = model_name or os.getenv("DOCX_EXPORT_NORMALIZER_MODEL", DEFAULT_MODEL_MINI)
        self.temperature = temperature
        self.token_usage_service = TokenUsageService()

    @staticmethod
    def _project_root() -> Path:
        return Path(__file__).resolve().parents[3]

    @classmethod
    def _template_path(cls) -> Path:
        return cls._project_root() / "templates_padrao" / "modelo_documento.docx"

    @staticmethod
    def _pick_style(style_names: list[str], candidates: list[str]) -> str | None:
        for candidate in candidates:
            if candidate in style_names:
                return candidate
        return None

    @staticmethod
    def _normalize_style_name(value: str) -> str:
        text = str(value or "").strip().lower()
        text = (
            text.replace("Г§", "c")
            .replace("ГЈ", "a")
            .replace("ГЎ", "a")
            .replace("Г ", "a")
            .replace("Гў", "a")
            .replace("Г©", "e")
            .replace("ГЄ", "e")
            .replace("Г­", "i")
            .replace("Гі", "o")
            .replace("Гґ", "o")
            .replace("Гµ", "o")
            .replace("Гє", "u")
        )
        return re.sub(r"\s+", " ", text)

    @classmethod
    def _pick_style_by_aliases(cls, style_names: list[str], aliases: list[str]) -> str | None:
        normalized_map = {cls._normalize_style_name(name): name for name in style_names}
        for alias in aliases:
            candidate = normalized_map.get(cls._normalize_style_name(alias))
            if candidate:
                return candidate
        return None

    @staticmethod
    def _strip_heading_number(text: str) -> str:
        value = str(text or "").strip()
        value = re.sub(r"^\d+(?:\.\d+)*\.?\s+", "", value)
        return value.strip()

    @staticmethod
    def _normalize_md_emphasis(text: str) -> str:
        value = str(text or "")
        value = value.replace("**", "")
        value = value.replace("__", "")
        return value

    @staticmethod
    def _unwrap_bold_heading_candidate(text: str) -> str:
        value = str(text or "").strip()
        if (value.startswith("**") and value.endswith("**") and len(value) > 4) or (
            value.startswith("__") and value.endswith("__") and len(value) > 4
        ):
            return value[2:-2].strip()
        return value

    @staticmethod
    def _clean_section_content(content: str) -> str:
        """Remove cabeГ§alhos duplicados no inГ­cio da seГ§ГЈo para preservar hierarquia."""
        text = str(content or "").strip()
        if not text:
            return ""

        text = re.sub(r"^#{1,6}\s+", "", text).strip()
        text = re.sub(r"^(?:[IVXLCDM]{1,8}|\d+(?:\.\d+)*)\.\s+", "", text, flags=re.IGNORECASE).strip()
        text = re.sub(
            r"^(?:DA|DO|DOS|DAS)?\s*"
            r"(INTRODUГ‡ГѓO|FATOS|FUNDAMENTOS|JURISPRUD[ГЉE]NCIA|PEDIDOS|CONCLUS[ГѓA]O)\s*[:\-]?\s*",
            "",
            text,
            flags=re.IGNORECASE,
        ).strip()
        return text

    @staticmethod
    def _clean_section_title(title: str) -> str:
        text = str(title or "").strip()
        text = re.sub(r"^#{1,6}\s+", "", text).strip()
        text = re.sub(r"^(?:[IVXLCDM]{1,8}|\d+(?:\.\d+)*)\.\s+", "", text, flags=re.IGNORECASE).strip()
        return text.upper()

    @classmethod
    def _is_plain_heading(cls, line: str) -> bool:
        value = str(line or "").strip()
        if not value or len(value) > 120:
            return False
        if value.startswith(("-", "*", ">", "|")):
            return False
        if value.startswith("RECURSO:"):
            return True
        if re.fullmatch(r"(?:[IVXLCDM]{1,8}|\d+(?:\.\d+)*)\.\s+.+", value):
            return True
        if value.isupper() and _PLAIN_HEADING_RE.fullmatch(value):
            return True
        return False

    # Detecta "3.1. HEADING CAPS Body em minГєsculas..." e insere \n entre eles.
    _HEADING_CONCAT_RE = re.compile(
        r'^(\d+(?:\.\d+)*\.\s+'
        r'(?:[A-ZГЃГ‰ГЌГ“ГљГЂГ‚ГЉГ”ГѓГ•Г‡0-9][A-ZГЃГ‰ГЌГ“ГљГЂГ‚ГЉГ”ГѓГ•Г‡0-9/,.()\- ]*?)'
        r')(?=\s+[A-ZГЃГ‰ГЌГ“ГљГЂГ‚ГЉГ”ГѓГ•Г‡][a-zГЎГ©Г­ГіГєГ ГўГЄГґГЈГµГ§])',
        re.MULTILINE,
    )

    def _split_concatenated_headings(self, text: str) -> str:
        """Insert \\n between a numbered all-caps heading and body text on the same line."""
        def _replace(m: re.Match) -> str:
            heading = m.group(1).rstrip()
            rest = text[m.end():m.end()]  # empty вЂ” lookahead doesn't consume
            return heading + "\n"

        result = []
        for line in text.splitlines(keepends=True):
            stripped = line.rstrip("\r\n")
            ending = line[len(stripped):]
            m = self._HEADING_CONCAT_RE.match(stripped)
            if m:
                heading = m.group(1).rstrip()
                body = stripped[m.end():].strip()
                if body:
                    result.append(heading + "\n" + body + ending)
                    continue
            result.append(line)
        return "".join(result)

    @staticmethod
    def _is_table_row(line: str) -> bool:
        stripped = str(line or "").strip()
        # Never treat ANEXO placeholders as table rows вЂ” they contain pipe chars
        # (e.g. {{ANEXO|id=X|arquivo=...|titulo=...|url=...}}) that would trigger
        # the 3-pipe heuristic below, corrupting the table accumulation buffer.
        if stripped.startswith("{{ANEXO|") and stripped.endswith("}}"):
            return False
        # Standard markdown: | cell | cell |
        if stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2:
            return True
        # Markdown without boundary pipes: cell | cell | cell (3+ pipes, not a sentence)
        if not stripped.startswith("|") and stripped.count("|") >= 3 and not stripped.endswith((".", "?", "!")):
            return True
        # Tab-separated (LLM sometimes generates TSV-style): 2+ tabs
        if "\t" in stripped and stripped.count("\t") >= 2:
            return True
        return False

    @staticmethod
    def _parse_table_cells(line: str) -> list[str]:
        stripped = str(line or "").strip()
        # Tab-separated
        if "\t" in stripped and stripped.count("\t") >= 1:
            return [cell.strip() for cell in stripped.split("\t")]
        # Pipe-separated
        stripped = stripped.strip("|")
        return [cell.strip() for cell in stripped.split("|")]

    @staticmethod
    def _is_separator_row(line: str) -> bool:
        stripped = str(line or "").strip()
        return bool(re.fullmatch(r"\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?", stripped))

    @staticmethod
    def _set_cell_shading(cell, fill_hex: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        for child in list(tc_pr):
            if child.tag == qn("w:shd"):
                tc_pr.remove(child)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_hex)
        tc_pr.append(shd)

    def _apply_table_visual_fallback(self, table, total_rows: int) -> None:
        # Fallback visual baseado no modelo legado modelo_acidente_trajeto.docx.
        if total_rows <= 0:
            return

        for col_index in range(len(table.rows[0].cells)):
            self._set_cell_shading(table.rows[0].cells[col_index], "D5DCE4")

        for row_index in range(1, total_rows):
            for col_index in range(len(table.rows[row_index].cells)):
                self._set_cell_shading(table.rows[row_index].cells[col_index], "FFFFFF")

    def _flush_markdown_table(self, doc, table_rows: list[list[str]], body_style: str | None) -> None:
        if not table_rows:
            return

        max_cols = max(len(row) for row in table_rows)
        if max_cols == 0:
            return

        table = doc.add_table(rows=len(table_rows), cols=max_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        section = doc.sections[-1] if doc.sections else None
        if section:
            table_width = section.page_width - section.left_margin - section.right_margin
            col_width = int(table_width / max_cols)
            table.autofit = False
            for row in table.rows:
                for cell in row.cells:
                    cell.width = col_width

        table_style = self._pick_style(
            [style.name for style in doc.styles if style.type == 3],
            [
                "Table Grid",
                "Grid Table Light",
                "Grid Table 1 Light",
                "Grid Table 1 Light Accent 1",
                "Grid Table 1 Light Accent 3",
                "Normal Table",
            ],
        )
        if table_style:
            table.style = table_style

        for row_index, row_values in enumerate(table_rows):
            for col_index in range(max_cols):
                cell = table.rows[row_index].cells[col_index]
                cell_text = row_values[col_index] if col_index < len(row_values) else ""
                cell.text = self._normalize_md_emphasis(cell_text)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        run.font.name = "Avenir Next LT Pro"
                        run.font.size = Pt(7)
                        if row_index == 0:
                            run.bold = True
                        else:
                            run.bold = False

        self._apply_table_visual_fallback(table, len(table_rows))

        doc.add_paragraph("")

    def _append_plain_line(self, doc, line: str, body_style: str | None) -> None:
        line = self._normalize_md_emphasis(line)
        if body_style:
            p = doc.add_paragraph(line, style=body_style)
        else:
            p = doc.add_paragraph(line)
            run = p.runs[0] if p.runs else p.add_run(line)
            run.font.name = "Segoe UI"
            run.font.size = Pt(11)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _append_rich_line(self, doc, line: str, body_style: str | None) -> None:
        paragraph = doc.add_paragraph(style=body_style) if body_style else doc.add_paragraph()
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        text = str(line or "")
        cursor = 0
        has_content = False

        for match in _BOLD_MARKDOWN_RE.finditer(text):
            before = text[cursor:match.start()]
            if before:
                run = paragraph.add_run(before)
                run.font.name = "Segoe UI"
                run.font.size = Pt(11)
                has_content = True

            bold_text = match.group(2) or match.group(3) or ""
            if bold_text:
                run = paragraph.add_run(bold_text)
                run.font.name = "Segoe UI"
                run.font.size = Pt(11)
                run.bold = True
                has_content = True

            cursor = match.end()

        remaining = text[cursor:]
        if remaining or not has_content:
            run = paragraph.add_run(remaining if remaining else text)
            run.font.name = "Segoe UI"
            run.font.size = Pt(11)

        return paragraph

    def _append_heading(self, doc, heading: str, level: int, title_style: str | None, subtitle_style: str | None):
        # Preserva numeraГ§ГЈo jГЎ existente no heading para evitar "reset" do Г­ndice.
        clean_heading = self._normalize_md_emphasis(heading).strip()

        paragraph_style = title_style if level == 1 else subtitle_style
        if paragraph_style:
            paragraph = doc.add_paragraph(clean_heading, style=paragraph_style)
            if not paragraph.runs:
                paragraph.add_run(clean_heading)
        else:
            # Evita estilos numerados do template que podem reiniciar a contagem.
            paragraph = doc.add_paragraph(clean_heading)
            target_size = 16 if level == 1 else 13

            if not paragraph.runs:
                paragraph.add_run(clean_heading)

            for run in paragraph.runs:
                run.font.name = "Segoe UI"
                run.font.size = Pt(target_size)
                run.bold = True

            paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 8)
            paragraph.paragraph_format.space_after = Pt(6)

        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return paragraph

    @staticmethod
    def _should_page_break_before_heading(heading: str) -> bool:
        normalized = str(heading or "").upper()
        normalized = (
            normalized.replace("Г‡", "C")
            .replace("Гѓ", "A")
            .replace("ГЃ", "A")
            .replace("ГЂ", "A")
            .replace("Г‚", "A")
            .replace("Г‰", "E")
            .replace("ГЉ", "E")
            .replace("ГЌ", "I")
            .replace("Г“", "O")
            .replace("Г”", "O")
            .replace("Г•", "O")
            .replace("Гљ", "U")
        )
        return ("PEDIDOS" in normalized) or ("PRELIMINAR" in normalized)

    def _sanitize_redundant_section_leads(self, document_text: str) -> str:
        """Remove repetiГ§ГЈo de rГіtulo de seГ§ГЈo logo apГіs o heading."""
        lines = str(document_text or "").splitlines()
        if not lines:
            return ""

        section_tokens = (
            "INTRODUCAO",
            "FATOS",
            "FUNDAMENTOS",
            "JURISPRUDENCIA",
            "PEDIDOS",
            "CONCLUSAO",
        )

        def normalize_heading_key(raw: str) -> str:
            value = self._normalize_md_emphasis(self._strip_heading_number(raw or "")).upper().strip()
            value = (
                value.replace("Г‡", "C")
                .replace("Гѓ", "A")
                .replace("ГЃ", "A")
                .replace("ГЂ", "A")
                .replace("Г‚", "A")
                .replace("Г‰", "E")
                .replace("ГЉ", "E")
                .replace("ГЌ", "I")
                .replace("Г“", "O")
                .replace("Г”", "O")
                .replace("Г•", "O")
                .replace("Гљ", "U")
            )
            value = re.sub(r"\s+", " ", value)
            return value

        for idx, raw_line in enumerate(lines):
            line = raw_line.strip()
            if not line:
                continue

            heading_text = None
            if line.startswith("#"):
                heading_text = re.sub(r"^#{1,6}\s*", "", line).strip()
            elif self._is_plain_heading(line):
                heading_text = line

            if not heading_text:
                continue

            normalized = normalize_heading_key(heading_text)
            if not any(token in normalized for token in section_tokens):
                continue

            next_idx = idx + 1
            while next_idx < len(lines) and not lines[next_idx].strip():
                next_idx += 1
            if next_idx >= len(lines):
                continue

            cleaned_next = self._clean_section_content(lines[next_idx])
            if cleaned_next and cleaned_next != lines[next_idx].strip():
                lines[next_idx] = cleaned_next

        return "\n".join(lines)

    def _normalize_for_docx(self, document_title: str, document_text: str, law_firm_id: Optional[int] = None) -> str:
        """Normaliza o texto com LLM mini. Em falha, retorna texto original."""
        source_text = str(document_text or "").strip()
        if not source_text:
            return ""

        llm = ChatOpenAI(
            model=self.model_name,
            temperature=self.temperature,
            max_tokens=8_000,
        ).with_structured_output(NormalizedDocxText, include_raw=True)

        user_prompt = (
            "Normalize o texto abaixo para exportaГ§ГЈo em DOCX.\n"
            "Regras:\n"
            "- Preserve todo o conteГєdo jurГ­dico.\n"
            "- NГЈo resuma, nГЈo corte, nГЈo invente.\n"
            "- Pode manter tГ­tulos em markdown (#, ##, ###) e separadores (---).\n"
            "- Remova apenas ruГ­do de formataГ§ГЈo e quebras excessivas.\n\n"
            f"TГ­tulo do documento: {document_title}\n\n"
            "=== TEXTO ORIGINAL ===\n"
            f"{source_text}"
        )

        messages = [
            {"role": "system", "content": _SYSTEM_PROMPT},
            {"role": "user", "content": user_prompt},
        ]

        started = time.perf_counter()
        raw_message = None
        try:
            output = llm.invoke(messages)
            raw_message = output.get("raw") if isinstance(output, dict) else None
            parsed = output.get("parsed") if isinstance(output, dict) else None
            normalized = str(getattr(parsed, "normalized_text", "") or "").strip()
            return normalized or source_text
        except Exception:
            return source_text
        finally:
            elapsed_ms = int((time.perf_counter() - started) * 1000)
            try:
                self.token_usage_service.capture_and_store(
                    response_payload={"messages": [raw_message]} if raw_message is not None else None,
                    agent_name="OfficeDocxExportAgent",
                    action_name="normalize_for_docx",
                    print_prefix="[OfficeDocxExportAgent][TokenUsage]",
                    model_name=self.model_name,
                    model_provider="openai",
                    law_firm_id=law_firm_id,
                    latency_ms=elapsed_ms,
                    status="success" if raw_message is not None else "error",
                    metadata_payload={"document_title": document_title[:120]},
                )
            except Exception:
                pass

    def _render_docx_with_template(
        self,
        document_title: str,
        document_text: str,
        include_document_title: bool = True,
    ) -> BytesIO:
        template_path = self._template_path()

        if template_path.exists():
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                temp_path = Path(tmp.name)
            try:
                shutil.copy2(str(template_path), str(temp_path))
                doc = Document(str(temp_path))
            finally:
                if temp_path.exists():
                    temp_path.unlink(missing_ok=True)
        else:
            doc = Document()
            for section in doc.sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1.25)
                section.right_margin = Inches(1.25)

        style_names = [style.name for style in doc.styles if style.type == 1]
        title_style = self._pick_style(style_names, ["TГ­tulo", "Titulo", "Heading 1", "TГ­tulo 1"])
        subtitle_style = self._pick_style(style_names, ["TГ­tulo 2", "Titulo 2", "Heading 2", "SubtГ­tulo"])

        topic_style = self._pick_style_by_aliases(
            style_names,
            [
                "TГіpico tese",
                "Topico tese",
                "Topico Tese",
                "TГіpico Tese",
            ],
        )
        subtopic_style = self._pick_style_by_aliases(
            style_names,
            [
                "SubtГіpico tese",
                "Subtopico tese",
                "SubtГіpico Tese",
                "Subtopico Tese",
                "Sub tГіpico tese",
                "Sub topico tese",
            ],
        )

        effective_heading_style = topic_style or title_style
        # Regra solicitada: usar o mesmo estilo de tГіpico tambГ©m nos subtГ­tulos.
        effective_subheading_style = effective_heading_style
        body_style = self._pick_style(style_names, ["Corpo de Texto", "Corpo de texto", "Normal", "Body Text"])

        for _ in range(len(doc.paragraphs)):
            paragraph = doc.paragraphs[0]
            paragraph._element.getparent().remove(paragraph._element)

        if include_document_title:
            if title_style:
                title_paragraph = doc.add_paragraph(str(document_title or "DOCUMENTO GERADO").upper(), style=title_style)
            else:
                title_paragraph = doc.add_paragraph()
                title_run = title_paragraph.add_run(str(document_title or "DOCUMENTO GERADO").upper())
                title_run.font.name = "Segoe UI"
                title_run.font.size = Pt(16)
                title_run.bold = True
            if not title_paragraph.runs:
                title_paragraph.add_run(str(document_title or "DOCUMENTO GERADO").upper())
            for run in title_paragraph.runs:
                run.font.name = "Segoe UI"
                run.font.size = Pt(16)
                run.bold = True
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        text = str(document_text or "")
        text = self._split_concatenated_headings(text)
        lines = text.splitlines()
        pending_table: list[list[str]] = []
        in_table = False

        def flush_table() -> None:
            nonlocal pending_table, in_table
            if pending_table:
                self._flush_markdown_table(doc, pending_table, body_style)
            pending_table = []
            in_table = False

        for raw_line in lines:
            line = raw_line.strip()
            if not line:
                flush_table()
                doc.add_paragraph("")
                continue

            bold_heading_candidate = self._unwrap_bold_heading_candidate(line)
            if bold_heading_candidate and self._is_plain_heading(bold_heading_candidate):
                flush_table()
                heading_level = 1 if bold_heading_candidate.upper().startswith("RECURSO:") else 2
                if self._should_page_break_before_heading(bold_heading_candidate) and doc.paragraphs:
                    doc.add_page_break()
                self._append_heading(
                    doc,
                    bold_heading_candidate,
                    level=heading_level,
                    title_style=effective_heading_style,
                    subtitle_style=effective_subheading_style,
                )
                continue

            if self._is_separator_row(line):
                continue

            if self._is_table_row(line):
                cells = self._parse_table_cells(line)
                if cells:
                    pending_table.append(cells)
                    in_table = True
                continue

            flush_table()

            if line.startswith("# "):
                heading_value = line[2:].strip()
                if self._should_page_break_before_heading(heading_value) and doc.paragraphs:
                    doc.add_page_break()
                self._append_heading(
                    doc,
                    heading_value,
                    level=1,
                    title_style=effective_heading_style,
                    subtitle_style=effective_subheading_style,
                )
                continue

            if line.startswith("## ") or line.startswith("### "):
                heading = line[3:].strip() if line.startswith("## ") else line[4:].strip()
                if self._should_page_break_before_heading(heading) and doc.paragraphs:
                    doc.add_page_break()
                self._append_heading(
                    doc,
                    heading,
                    level=2,
                    title_style=effective_heading_style,
                    subtitle_style=effective_subheading_style,
                )
                continue

            if self._is_plain_heading(line):
                heading_level = 1 if line.upper().startswith("RECURSO:") else 2
                if self._should_page_break_before_heading(line) and doc.paragraphs:
                    doc.add_page_break()
                self._append_heading(
                    doc,
                    line,
                    level=heading_level,
                    title_style=effective_heading_style,
                    subtitle_style=effective_subheading_style,
                )
                continue

            if line.startswith("---"):
                p = doc.add_paragraph("_" * 60)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                continue

            # Linha que Г© APENAS uma referГЄncia de anexo
            if re.fullmatch(r"\[Anexo:[^\]]+\]", line):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.name = "Segoe UI"
                run.font.size = Pt(10)
                run.italic = True
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                continue

            if "**" in line or "__" in line:
                self._append_rich_line(doc, line, body_style)
            else:
                self._append_plain_line(doc, line, body_style)

        flush_table()

        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output

    _ANEXO_RE = re.compile(r"\{\{ANEXO\|([^}]+)\}\}")

    def _strip_anexo_placeholders(self, text: str) -> str:
        """Converte {{ANEXO|...}} em referГЄncia textual legГ­vel para o DOCX."""
        def _replace(m: re.Match) -> str:
            params: dict[str, str] = {}
            for part in m.group(1).split("|"):
                if "=" in part:
                    k, v = part.split("=", 1)
                    params[k.strip()] = v.strip()
            titulo = params.get("titulo", "")
            arquivo = params.get("arquivo", "")
            if titulo and arquivo:
                return f"[Anexo: {titulo} ({arquivo})]"
            if titulo:
                return f"[Anexo: {titulo}]"
            if arquivo:
                return f"[Anexo: {arquivo}]"
            return ""

        return self._ANEXO_RE.sub(_replace, text)

    def export_generated_document(
        self,
        *,
        document_title: str,
        document_text: str,
        run_ai_normalization: bool = True,
        law_firm_id: Optional[int] = None,
        include_document_title: bool = True,
    ) -> BytesIO:
        """Exporta documento gerado para DOCX no padrГЈo do escritГіrio."""
        text = str(document_text or "")
        # Strip ANEXO placeholders before LLM normalization вЂ” the pipe-heavy
        # {{ANEXO|key=val|...}} syntax confuses the normalizer and corrupts nearby
        # markdown tables and heading structure.
        text = self._strip_anexo_placeholders(text)
        if run_ai_normalization:
            text = self._normalize_for_docx(
                document_title=document_title,
                document_text=text,
                law_firm_id=law_firm_id,
            )
        text = self._sanitize_redundant_section_leads(text)
        return self._render_docx_with_template(
            document_title=document_title,
            document_text=text,
            include_document_title=include_document_title,
        )

    def export_appeal_content(
        self,
        *,
        appeal_content: dict,
        run_ai_normalization: bool = True,
        law_firm_id: Optional[int] = None,
    ) -> BytesIO:
        """Converte payload estruturado de recurso em DOCX no padrГЈo do escritГіrio."""
        appeal_type = str(appeal_content.get("appeal_type") or "RECURSO JUDICIAL").strip()
        sections: list[str] = [f"# {appeal_type}"]

        if appeal_content.get("introduction"):
            sections.append("## I. INTRODUГ‡ГѓO")
            sections.append(self._clean_section_content(str(appeal_content["introduction"])))

        if appeal_content.get("facts"):
            sections.append("## II. DOS FATOS")
            sections.append(self._clean_section_content(str(appeal_content["facts"])))

        if appeal_content.get("grounds"):
            sections.append("## III. DOS FUNDAMENTOS")
            sections.append(self._clean_section_content(str(appeal_content["grounds"])))

        if appeal_content.get("jurisprudence"):
            sections.append("## IV. DA JURISPRUDГЉNCIA")
            sections.append(self._clean_section_content(str(appeal_content["jurisprudence"])))

        section_num = 5
        for section in appeal_content.get("additional_sections", []) or []:
            title = self._clean_section_title(str(section.get("title") or "SEГ‡ГѓO ADICIONAL"))
            content = self._clean_section_content(str(section.get("content") or ""))
            if content:
                sections.append(f"## {section_num}. {title}")
                sections.append(content)
                section_num += 1

        if appeal_content.get("requests"):
            sections.append(f"## {section_num}. DOS PEDIDOS")
            sections.append(self._clean_section_content(str(appeal_content["requests"])))

        if appeal_content.get("conclusion"):
            sections.append("## CONCLUSГѓO")
            sections.append(self._clean_section_content(str(appeal_content["conclusion"])))

        merged_text = "\n\n".join(part for part in sections if part)
        return self.export_generated_document(
            document_title=appeal_type,
            document_text=merged_text,
            run_ai_normalization=run_ai_normalization,
            law_firm_id=law_firm_id,
        )
