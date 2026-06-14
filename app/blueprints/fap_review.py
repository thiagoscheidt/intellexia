"""
Blueprint: FAP Review - Módulo de Revisão de Petição Inicial FAP

Rotas principais:
- /fap-review - Dashboard principal
- /fap-review/revision - Revisão de petições
- /fap-review/training - Gerenciamento de treinamento
- /fap-review/settings - Configurações do módulo
"""

import os
import json
import asyncio
import hashlib
import re
import unicodedata
from datetime import datetime
from io import BytesIO
from decimal import Decimal
from pathlib import Path

from flask import Blueprint, current_app, flash, jsonify, redirect, render_template, request, session, url_for, send_file
from werkzeug.utils import secure_filename
from sqlalchemy import and_, func, case

from app.models import (
    db, User, LawFirm,
    FapReviewPetition,
    FapReviewPromptVersion, FapReviewReferenceVersion, FapReviewSetting,
    FapReviewExecution, FapReviewIgnoredFinding, FapReviewAuditLog
)
from app.agents.fap_review import (
    FapPetitionReviewerAgent,
    FapTrainingEvolutionAgent,
    FapTrainingApplySubAgent,
)
from app.services.openrouter_models_service import fetch_openrouter_text_models_for_info
from app.utils.timezone import now_sp

# Document processing
try:
    from docling.document_converter import DocumentConverter
    HAS_DOCLING = True
except ImportError:
    HAS_DOCLING = False

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from openpyxl import load_workbook
except ImportError:
    load_workbook = None


fap_review_bp = Blueprint('fap_review', __name__, url_prefix='/fap-review')

# Configurações
ALLOWED_DOCUMENT_EXTENSIONS = {'pdf', 'doc', 'docx', 'txt'}
ALLOWED_AUXILIARY_EXTENSIONS = {'pdf', 'doc', 'docx', 'xls', 'xlsx', 'txt', 'png', 'jpg', 'jpeg'}
ALLOWED_BENEFITS_SPREADSHEET_EXTENSIONS = {'xlsx'}
MAX_UPLOAD_SIZE = 50 * 1024 * 1024  # 50MB
READ_ONLY_PROMPT_TYPES = {'revisor_output_format'}
READ_ONLY_REFERENCE_TYPES = {'project_instructions'}
PETITION_WORKFLOW_STATUSES = {
    'new': 'Nova',
    'in_review': 'Em revisão',
    'awaiting_adjustments': 'Aguardando ajustes',
    'ready_for_filing': 'Pronta para seguir',
    'filed': 'Processo iniciado',
    'archived': 'Arquivada',
}


def _get_file_extension(filename: str) -> str:
    """Obtém extensão normalizada do arquivo."""
    if not filename or '.' not in filename:
        return ''
    return f".{filename.rsplit('.', 1)[1].lower()}"


def allowed_file(filename: str, allowed_extensions: set) -> bool:
    """Verifica se arquivo tem extensão permitida"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions


def get_current_law_firm_id() -> int:
    """Obtém ID do escritório da sessão"""
    return session.get('law_firm_id')


def require_law_firm(f):
    """Decorator para garantir que há escritório na sessão"""
    from functools import wraps
    
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not get_current_law_firm_id():
            flash('Acesso negado', 'error')
            return redirect(url_for('auth.login'))
        return f(*args, **kwargs)
    
    return decorated_function


def require_admin_user(f):
    """Decorator para garantir que é admin"""
    from functools import wraps
    
    @wraps(f)
    def decorated_function(*args, **kwargs):
        user_id = session.get('user_id')
        if not user_id:
            return redirect(url_for('auth.login'))
        
        user = User.query.get(user_id)
        if not user or user.role != 'admin':
            flash('Acesso negado: privilégio de administrador necessário', 'error')
            return redirect(url_for('fap_review.index'))
        
        return f(*args, **kwargs)
    
    return decorated_function


def _get_fap_setting(law_firm_id: int) -> FapReviewSetting:
    """Obtém ou cria configuração padrão do FAP Review"""
    setting = FapReviewSetting.query.filter_by(law_firm_id=law_firm_id).first()
    
    if not setting:
        setting = FapReviewSetting(
            law_firm_id=law_firm_id,
            reviewer_model='gpt-4o-mini',
            training_model='gpt-4o-mini',
            reviewer_temperature=0.0,  # Temperature=0.0 com seed para determinismo garantido
            training_temperature=0.7,
            reviewer_enabled=True,
            training_enabled=True
        )
        db.session.add(setting)
        db.session.commit()
    
    return setting


def _log_audit(law_firm_id: int, action: str, entity_type: str, 
               entity_id: int = None, description: str = "", 
               old_value: str = "", new_value: str = ""):
    """Registra ação de auditoria"""
    user_id = session.get('user_id')
    
    log_entry = FapReviewAuditLog(
        law_firm_id=law_firm_id,
        user_id=user_id,
        action=action,
        entity_type=entity_type,
        entity_id=entity_id,
        change_description=description,
        old_value=old_value,
        new_value=new_value
    )
    db.session.add(log_entry)
    db.session.commit()


def _create_upload_directory(law_firm_id: int, subdir: str = "") -> Path:
    """Cria diretório de upload se não existir"""
    base_dir = Path('uploads/fap_review') / str(law_firm_id)
    if subdir:
        base_dir = base_dir / subdir
    base_dir.mkdir(parents=True, exist_ok=True)
    return base_dir


def _extract_text_from_document(filepath: str) -> str:
    """Extrai texto de um documento (PDF, DOCX ou TXT)."""
    filepath = Path(filepath)
    extension = filepath.suffix.lower()
    text = ""
    
    try:
        if extension == '.pdf':
            # Tentar Docling primeiro (melhor qualidade)
            if HAS_DOCLING:
                try:
                    converter = DocumentConverter()
                    doc_result = converter.convert(str(filepath))
                    text = doc_result.document.export_to_markdown()
                except Exception as e:
                    current_app.logger.warning(f"Docling falhou: {e}, tentando PyPDF2")
                    # Fallback para PyPDF2
                    if PyPDF2:
                        with open(filepath, 'rb') as f:
                            reader = PyPDF2.PdfReader(f)
                            for page in reader.pages:
                                text += page.extract_text() + "\n"
            elif PyPDF2:
                with open(filepath, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        text += page.extract_text() + "\n"
            else:
                raise ImportError("PyPDF2 não está instalado")
        
        elif extension == '.docx':
            if DocxDocument:
                try:
                    doc = DocxDocument(filepath)
                    for para in doc.paragraphs:
                        text += para.text + "\n"
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                text += cell.text + " | "
                            text += "\n"
                except Exception as e:
                    raise ValueError(f"Erro ao ler DOCX: {e}")
            else:
                raise ImportError("python-docx não está instalado")

        elif extension == '.doc':
            raise ValueError("Arquivos .doc não são suportados no FAP Review. Envie em PDF ou DOCX.")
        
        elif extension == '.txt':
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()
        
        else:
            raise ValueError(f"Tipo de arquivo não suportado: {extension}")
        
        # Limpar espaços em branco excessivos
        text = '\n'.join([line.strip() for line in text.split('\n') if line.strip()])
        
        return text
    
    except Exception as e:
        current_app.logger.error(f"Erro ao extrair texto: {e}")
        raise


def _normalize_spreadsheet_header(value: object) -> str:
    """Normaliza cabeçalhos de planilha para busca resiliente."""
    normalized = unicodedata.normalize('NFKD', str(value or ''))
    ascii_text = normalized.encode('ascii', 'ignore').decode('ascii')
    return ' '.join(ascii_text.strip().lower().split())


def _format_benefit_number(value: object) -> str:
    """Formata número de benefício preservando apenas o valor legível."""
    if value is None:
        return ''

    if isinstance(value, bool):
        return ''

    if isinstance(value, int):
        return str(value)

    if isinstance(value, float) and value.is_integer():
        return str(int(value))

    return str(value).strip()


def _normalize_benefit_number(value: object) -> str:
    """Normaliza número de benefício para comparação textual."""
    return re.sub(r'\D+', '', _format_benefit_number(value))


def _document_mentions_benefit_number(document_text: str, normalized_benefit_number: str) -> bool:
    """Verifica se o documento menciona o benefício com tolerância a pontuação e separadores."""
    if not document_text or not normalized_benefit_number:
        return False

    # Aceita formatos como 1234567890, 123.456.789-0, 123 456 789 0 ou variações com separadores.
    separator_pattern = r'[\s\.\-\/_;,:\(\)\[\]]*'
    digit_pattern = separator_pattern.join(re.escape(digit) for digit in normalized_benefit_number)
    pattern = rf'(?<!\d){digit_pattern}(?!\d)'
    return re.search(pattern, document_text) is not None


def _parse_benefits_spreadsheet(filepath: str) -> list[dict[str, str]]:
    """Lê a planilha de benefícios e retorna apenas linhas com tese preenchida."""
    if not load_workbook:
        raise ImportError('openpyxl não está instalado')

    workbook = load_workbook(filename=filepath, read_only=True, data_only=True)
    worksheet = workbook.active

    try:
        header_row = next(worksheet.iter_rows(min_row=1, max_row=1, values_only=True), None)
        if not header_row:
            raise ValueError('A planilha de benefícios está vazia.')

        header_map = {
            _normalize_spreadsheet_header(value): index
            for index, value in enumerate(header_row)
        }
        benefit_idx = header_map.get('numero do beneficio')
        thesis_idx = header_map.get('teses')

        if benefit_idx is None or thesis_idx is None:
            raise ValueError(
                'A planilha deve conter as colunas "Número do Benefício" e "TESES".'
            )

        rows: list[dict[str, str]] = []
        for row in worksheet.iter_rows(min_row=2, values_only=True):
            thesis = ' '.join(str(row[thesis_idx] or '').strip().split()) if thesis_idx < len(row) else ''
            if not thesis:
                continue

            benefit_number = _format_benefit_number(row[benefit_idx] if benefit_idx < len(row) else '')
            normalized_benefit_number = _normalize_benefit_number(benefit_number)
            if not normalized_benefit_number:
                continue

            rows.append({
                'benefit_number': benefit_number,
                'benefit_number_normalized': normalized_benefit_number,
                'thesis': thesis,
            })

        return rows
    finally:
        workbook.close()


def _build_benefits_spreadsheet_review(
    spreadsheet_path: str,
    spreadsheet_filename: str,
    document_text: str,
    checked_document_label: str,
) -> dict:
    """Gera checagem determinística dos benefícios da planilha contra o documento."""
    rows = _parse_benefits_spreadsheet(spreadsheet_path)

    items: list[dict[str, object]] = []
    found_count = 0

    for row in rows:
        found_in_document = _document_mentions_benefit_number(
            document_text=document_text,
            normalized_benefit_number=row['benefit_number_normalized'],
        )
        if found_in_document:
            found_count += 1

        items.append({
            'benefit_number': row['benefit_number'],
            'thesis': row['thesis'],
            'found_in_document': found_in_document,
        })

    return {
        'source_filename': spreadsheet_filename,
        'checked_document_label': checked_document_label,
        'total_rows_with_theses': len(items),
        'found_count': found_count,
        'missing_count': max(len(items) - found_count, 0),
        'items': items,
    }


def _normalize_finding_field(value: object) -> str:
    """Normaliza campo textual de achado para fingerprint estável."""
    return ' '.join(str(value or '').strip().split()).lower()


def _build_finding_fingerprint(finding: dict | None) -> str:
    """Gera fingerprint estável de um achado para persistir feedback do usuário."""
    if not isinstance(finding, dict):
        return ''

    payload = {
        'category': _normalize_finding_field(finding.get('category')),
        'severity': _normalize_finding_field(finding.get('severity')),
        'description': _normalize_finding_field(finding.get('description')),
        'location': _normalize_finding_field(finding.get('location')),
        'correction': _normalize_finding_field(finding.get('correction')),
        'manual_reference': _normalize_finding_field(finding.get('manual_reference')),
    }
    serialized = json.dumps(payload, sort_keys=True, ensure_ascii=False)
    return hashlib.sha256(serialized.encode('utf-8')).hexdigest()


def _get_ignored_finding_fingerprints(law_firm_id: int, law_firm_document_identifier: str) -> set[str]:
    """Retorna fingerprints de achados marcados como não úteis para o documento."""
    if not law_firm_document_identifier:
        return set()

    rows = FapReviewIgnoredFinding.query.filter_by(
        law_firm_id=law_firm_id,
        law_firm_document_identifier=law_firm_document_identifier,
    ).with_entities(FapReviewIgnoredFinding.finding_fingerprint).all()
    return {str(row[0]) for row in rows if row and row[0]}


def _build_prior_attention_points(
    law_firm_id: int,
    petition_id: int | None,
    law_firm_document_identifier: str,
    current_execution_id: int,
) -> str | None:
    """Resume os pontos de atenção da última revisão concluída do mesmo documento."""
    if not petition_id and not law_firm_document_identifier:
        return None

    previous_execution_query = FapReviewExecution.query.filter(
        FapReviewExecution.law_firm_id == law_firm_id,
        FapReviewExecution.execution_type == 'revision',
        FapReviewExecution.status == 'completed',
        FapReviewExecution.id != current_execution_id,
        FapReviewExecution.result_json.isnot(None),
    )

    if petition_id:
        previous_execution_query = previous_execution_query.filter(
            FapReviewExecution.petition_id == petition_id,
        )
    else:
        previous_execution_query = previous_execution_query.filter(
            FapReviewExecution.law_firm_document_identifier == law_firm_document_identifier,
        )

    previous_execution = previous_execution_query.order_by(
        FapReviewExecution.completed_at.desc(),
        FapReviewExecution.id.desc(),
    ).first()

    if not previous_execution or not previous_execution.result_json:
        return None

    try:
        result_payload = json.loads(previous_execution.result_json)
    except (TypeError, json.JSONDecodeError):
        current_app.logger.warning(
            'Falha ao interpretar o histórico da execução %s para o identificador %s',
            previous_execution.id,
            law_firm_document_identifier,
        )
        return None

    ignored_fingerprints = _get_ignored_finding_fingerprints(
        law_firm_id=law_firm_id,
        law_firm_document_identifier=law_firm_document_identifier,
    )
    lines: list[str] = []

    for finding in result_payload.get('findings') or []:
        finding_fingerprint = _build_finding_fingerprint(finding)
        if finding_fingerprint and finding_fingerprint in ignored_fingerprints:
            continue

        description = str(finding.get('description') or '').strip()
        if not description:
            continue

        severity = str(finding.get('severity') or 'ATENÇÃO').strip().upper()
        location = str(finding.get('location') or '').strip()
        correction = str(finding.get('correction') or '').strip()
        manual_reference = str(finding.get('manual_reference') or '').strip()

        parts = [f'- [{severity}] {description}']
        if location:
            parts.append(f'Local: {location}')
        if correction:
            parts.append(f'Correção esperada: {correction}')
        if manual_reference:
            parts.append(f'Referência manual: {manual_reference}')
        lines.append(' | '.join(parts))

    for missing_document in result_payload.get('missing_documents') or []:
        document_type = str(missing_document.get('document_type') or '').strip()
        if not document_type:
            continue

        thesis = str(missing_document.get('thesis') or '').strip()
        manual_reference = str(missing_document.get('manual_reference') or '').strip()
        parts = [f'- [DOCUMENTO] {document_type}']
        if thesis:
            parts.append(f'Tese relacionada: {thesis}')
        if manual_reference:
            parts.append(f'Referência manual: {manual_reference}')
        lines.append(' | '.join(parts))

    if not lines:
        return (
            '__NO_ACTIVE_PRIOR_ATTENTION_POINTS__\n'
            'Os pontos de atenção anteriores deste identificador foram marcados como não úteis '
            'e não devem ser cobrados novamente nesta revisão.'
        )

    return '\n'.join(lines)


def _execute_reviewer_agent(execution_id: int, law_firm_id: int, petition_file_path: str,
                           compared_file_path: str = None,
                           benefits_spreadsheet: dict | None = None) -> dict:
    """Executa o agente revisor e armazena resultado"""
    try:
        execution = FapReviewExecution.query.get(execution_id)
        if not execution:
            raise ValueError(f"Execução {execution_id} não encontrada")
        
        # Carregar configurações
        setting = _get_fap_setting(law_firm_id)
        openai_api_key = os.getenv('OPENAI_API_KEY')
        if not openai_api_key:
            raise ValueError("OPENAI_API_KEY não configurada")
        
        # Carregar referências ativas
        manual = FapReviewReferenceVersion.query.filter_by(
            law_firm_id=law_firm_id,
            reference_type='manual_fap',
            is_active=True
        ).first()
        
        cases = FapReviewReferenceVersion.query.filter_by(
            law_firm_id=law_firm_id,
            reference_type='casos_referencia',
            is_active=True
        ).first()
        
        project_instructions = FapReviewReferenceVersion.query.filter_by(
            law_firm_id=law_firm_id,
            reference_type='project_instructions',
            is_active=True
        ).first()

        # Carregar prompts ativos do revisor
        reviewer_identity_prompt = FapReviewPromptVersion.query.filter_by(
            law_firm_id=law_firm_id,
            prompt_type='revisor_identity',
            is_active=True
        ).first()

        reviewer_rules_prompt = FapReviewPromptVersion.query.filter_by(
            law_firm_id=law_firm_id,
            prompt_type='revisor_rules',
            is_active=True
        ).first()

        reviewer_output_format_prompt = FapReviewPromptVersion.query.filter_by(
            law_firm_id=law_firm_id,
            prompt_type='revisor_output_format',
            is_active=True
        ).first()
        
        # Instanciar agente revisor
        agent = FapPetitionReviewerAgent(
            openai_api_key=openai_api_key,
            model=setting.reviewer_model,
            temperature=setting.reviewer_temperature
        )
        
        # Carregar referências no agente
        agent.load_reference_documents(
            manual_md=manual.content if manual else "",
            cases_md=cases.content if cases else "",
            project_instructions_md=project_instructions.content if project_instructions else ""
        )
        
        # Executar análise (async)
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        
        try:
            petition_extension = Path(petition_file_path).suffix.lower()
            petition_text = _extract_text_from_document(petition_file_path) if petition_extension in {'.docx', '.txt'} else None
            prior_attention_points = _build_prior_attention_points(
                law_firm_id=law_firm_id,
                petition_id=execution.petition_id,
                law_firm_document_identifier=execution.law_firm_document_identifier or '',
                current_execution_id=execution.id,
            )

            compared_text = None
            if compared_file_path:
                compared_extension = Path(compared_file_path).suffix.lower()
                compared_text = _extract_text_from_document(compared_file_path) if compared_extension in {'.docx', '.txt'} else None

            if compared_file_path and execution.comparative_analysis:
                # Análise comparativa
                result = loop.run_until_complete(
                    agent.review_petition_comparative(
                        original_petition_file_path=petition_file_path,
                        revised_petition_file_path=compared_file_path,
                        original_petition_text=petition_text,
                        revised_petition_text=compared_text,
                        prior_attention_points=prior_attention_points,
                        reviewer_identity=reviewer_identity_prompt.content if reviewer_identity_prompt else "",
                        reviewer_rules=reviewer_rules_prompt.content if reviewer_rules_prompt else "",
                        reviewer_output_format=reviewer_output_format_prompt.content if reviewer_output_format_prompt else "",
                        execution_id=execution.id,
                        user_id=execution.user_id,
                        law_firm_id=law_firm_id,
                    )
                )
            else:
                # Análise simples
                result = loop.run_until_complete(
                    agent.review_petition_single_version(
                        petition_file_path=petition_file_path,
                        petition_text=petition_text,
                        prior_attention_points=prior_attention_points,
                        reviewer_identity=reviewer_identity_prompt.content if reviewer_identity_prompt else "",
                        reviewer_rules=reviewer_rules_prompt.content if reviewer_rules_prompt else "",
                        reviewer_output_format=reviewer_output_format_prompt.content if reviewer_output_format_prompt else "",
                        execution_id=execution.id,
                        user_id=execution.user_id,
                        law_firm_id=law_firm_id,
                    )
                )

            result_payload = result.model_dump(mode='json')

            if benefits_spreadsheet and benefits_spreadsheet.get('path'):
                target_document_path = compared_file_path if compared_file_path and execution.comparative_analysis else petition_file_path
                target_document_text = compared_text if compared_file_path and execution.comparative_analysis else petition_text
                if not target_document_text:
                    target_document_text = _extract_text_from_document(target_document_path)

                checked_document_label = (
                    'Documento revisado'
                    if compared_file_path and execution.comparative_analysis
                    else (execution.main_document_filename or 'Documento principal')
                )

                try:
                    result_payload['benefits_spreadsheet_review'] = _build_benefits_spreadsheet_review(
                        spreadsheet_path=str(benefits_spreadsheet['path']),
                        spreadsheet_filename=str(benefits_spreadsheet.get('name') or Path(benefits_spreadsheet['path']).name),
                        document_text=target_document_text,
                        checked_document_label=checked_document_label,
                    )
                except Exception as spreadsheet_error:
                    current_app.logger.warning(
                        'Falha ao processar planilha de beneficios da execucao %s: %s',
                        execution_id,
                        spreadsheet_error,
                    )
                    result_payload['benefits_spreadsheet_review'] = {
                        'source_filename': str(benefits_spreadsheet.get('name') or Path(benefits_spreadsheet['path']).name),
                        'checked_document_label': checked_document_label,
                        'error': str(spreadsheet_error),
                        'items': [],
                        'total_rows_with_theses': 0,
                        'found_count': 0,
                        'missing_count': 0,
                    }
            
            # Armazenar resultado
            execution.result_json = json.dumps(result_payload, ensure_ascii=False, indent=2)
            execution.status = 'completed'
            execution.completed_at = datetime.utcnow()
            
            # Se o resultado tem tokens, armazenar
            if hasattr(result, 'tokens_used'):
                execution.tokens_used = result.tokens_used
            if hasattr(result, 'cost_usd'):
                execution.cost_usd = Decimal(str(result.cost_usd))

            _sync_petition_after_revision(execution)
            
            db.session.commit()
            
            # Log de auditoria
            _log_audit(law_firm_id, 'revision_completed', 'execution', execution_id,
                      'Revisão concluída com sucesso')
            
            return {'success': True, 'execution_id': execution_id}
        
        finally:
            loop.close()
    
    except Exception as e:
        current_app.logger.error(f"Erro ao executar agente revisor: {e}")
        
        # Atualizar status para falha
        try:
            execution = FapReviewExecution.query.get(execution_id)
            if execution:
                execution.status = 'failed'
                execution.error_message = str(e)
                execution.completed_at = datetime.utcnow()
                _sync_petition_after_revision(execution)
                db.session.commit()
                
                _log_audit(law_firm_id, 'revision_failed', 'execution', execution_id,
                          f'Erro: {str(e)[:200]}')
        except Exception as inner_e:
            current_app.logger.error(f"Erro ao registrar falha: {inner_e}")
        
        raise


def _get_active_reference(law_firm_id: int, reference_type: str) -> FapReviewReferenceVersion | None:
    """Obtém a versão ativa mais recente de uma referência."""
    return FapReviewReferenceVersion.query.filter_by(
        law_firm_id=law_firm_id,
        reference_type=reference_type,
        is_active=True,
    ).order_by(FapReviewReferenceVersion.version_number.desc()).first()


def _append_reference_version(
    law_firm_id: int,
    user_id: int,
    reference_type: str,
    new_content: str,
    activate: bool = True,
) -> FapReviewReferenceVersion:
    """Cria uma nova versão de referência, opcionalmente ativando-a."""
    latest = FapReviewReferenceVersion.query.filter_by(
        law_firm_id=law_firm_id,
        reference_type=reference_type,
    ).order_by(FapReviewReferenceVersion.version_number.desc()).first()

    next_version = (latest.version_number + 1) if latest else 1

    if activate:
        FapReviewReferenceVersion.query.filter_by(
            law_firm_id=law_firm_id,
            reference_type=reference_type,
            is_active=True,
        ).update({'is_active': False})

    version = FapReviewReferenceVersion(
        law_firm_id=law_firm_id,
        version_number=next_version,
        reference_type=reference_type,
        content=new_content,
        is_active=activate,
        created_by_id=user_id,
    )
    db.session.add(version)
    return version


def _build_petition_title(raw_title: str, fallback_filename: str = '', fallback_identifier: str = '') -> str:
    """Monta um título amigável para a petição."""
    title = ' '.join(str(raw_title or '').strip().split())
    if title:
        return title

    filename = Path(str(fallback_filename or '')).stem.strip()
    if filename:
        return filename

    identifier = str(fallback_identifier or '').strip()
    return identifier or 'Petição sem título'


def _derive_petition_workflow_status(execution_status: str) -> str:
    """Traduz o status da revisão para o status agregado da petição."""
    if execution_status in {'pending', 'processing'}:
        return 'in_review'
    if execution_status == 'completed':
        return 'ready_for_filing'
    if execution_status == 'failed':
        return 'awaiting_adjustments'
    return 'new'


def _sync_petition_after_revision(execution: FapReviewExecution) -> None:
    """Atualiza a visão agregada da petição depois de uma revisão."""
    petition = execution.petition
    if not petition or execution.execution_type != 'revision':
        return

    if execution.revision_number is None:
        execution.revision_number = FapReviewExecution.query.filter_by(
            petition_id=petition.id,
            execution_type='revision',
        ).count()

    petition.latest_revision_id = execution.id
    petition.revision_count = max(petition.revision_count or 0, execution.revision_number or 0)
    petition.last_reviewed_at = execution.completed_at or execution.updated_at or execution.created_at

    if petition.workflow_status not in {'filed', 'archived'}:
        petition.workflow_status = _derive_petition_workflow_status(execution.status)

    petition.updated_at = datetime.utcnow()


def _build_petition_status_badge(workflow_status: str) -> dict[str, str]:
    """Retorna classes e label para o status da petição na UI."""
    mapping = {
        'new': {'class': 'secondary', 'icon': 'bi bi-plus-circle'},
        'in_review': {'class': 'warning', 'icon': 'bi bi-hourglass-split'},
        'awaiting_adjustments': {'class': 'danger', 'icon': 'bi bi-pencil-square'},
        'ready_for_filing': {'class': 'success', 'icon': 'bi bi-check-circle'},
        'filed': {'class': 'primary', 'icon': 'bi bi-send-check'},
        'archived': {'class': 'dark', 'icon': 'bi bi-archive'},
    }
    status_key = workflow_status if workflow_status in mapping else 'new'
    return {
        'label': PETITION_WORKFLOW_STATUSES.get(status_key, 'Nova'),
        **mapping[status_key],
    }


# ═══════════════════════════════════════════════════════════════════════════════
# ROTAS PRINCIPAIS
# ═══════════════════════════════════════════════════════════════════════════════


@fap_review_bp.route('/')
@require_law_firm
def index():
    """Dashboard principal do módulo"""
    law_firm_id = get_current_law_firm_id()

    total_petitions = FapReviewPetition.query.filter_by(law_firm_id=law_firm_id).count()
    ready_petitions = FapReviewPetition.query.filter_by(
        law_firm_id=law_firm_id,
        workflow_status='ready_for_filing',
    ).count()
    in_review_petitions = FapReviewPetition.query.filter(
        FapReviewPetition.law_firm_id == law_firm_id,
        FapReviewPetition.workflow_status.in_(['new', 'in_review']),
    ).count()
    awaiting_adjustments_petitions = FapReviewPetition.query.filter_by(
        law_firm_id=law_firm_id,
        workflow_status='awaiting_adjustments',
    ).count()
    total_revisions = FapReviewExecution.query.filter_by(
        law_firm_id=law_firm_id,
        execution_type='revision',
    ).count()

    _priority_order = case(
        (FapReviewPetition.workflow_status == 'awaiting_adjustments', 0),
        (FapReviewPetition.workflow_status.in_(['new', 'in_review']), 1),
        (FapReviewPetition.workflow_status == 'ready_for_filing', 2),
        else_=3
    )

    petitions = FapReviewPetition.query.filter_by(
        law_firm_id=law_firm_id,
    ).order_by(
        _priority_order,
        FapReviewPetition.updated_at.desc(),
        FapReviewPetition.id.desc(),
    ).limit(20).all()

    petition_rows = [
        {
            'petition': petition,
            'latest_revision': petition.latest_revision,
            'status_badge': _build_petition_status_badge(petition.workflow_status),
        }
        for petition in petitions
    ]

    return render_template('fap_review/index.html',
                          total_petitions=total_petitions,
                          ready_petitions=ready_petitions,
                          in_review_petitions=in_review_petitions,
                          awaiting_adjustments_petitions=awaiting_adjustments_petitions,
                          total_revisions=total_revisions,
                          petition_rows=petition_rows)


@fap_review_bp.route('/revision', methods=['GET', 'POST'])
@require_law_firm
def revision():
    """Página de revisão de petições"""
    law_firm_id = get_current_law_firm_id()
    user_id = session.get('user_id')
    
    if request.method == 'POST':
        try:
            # Validar upload de arquivo principal
            if 'main_document' not in request.files:
                return jsonify({'error': 'Documento principal não fornecido'}), 400
            
            main_file = request.files['main_document']
            if main_file.filename == '':
                return jsonify({'error': 'Arquivo vazio'}), 400

            main_extension = _get_file_extension(main_file.filename)
            if main_extension == '.doc':
                return jsonify({'error': 'Arquivos .doc não são suportados no FAP Review. Envie em PDF ou DOCX.'}), 400
            
            if not allowed_file(main_file.filename, ALLOWED_DOCUMENT_EXTENSIONS):
                return jsonify({'error': 'Tipo de arquivo não permitido'}), 400

            petition_id = request.form.get('petition_id', type=int)
            petition_title = str(request.form.get('petition_title', '')).strip()
            petition = None
            law_firm_document_identifier = str(
                request.form.get('law_firm_document_identifier', '')
            ).strip()

            if petition_id:
                petition = FapReviewPetition.query.filter_by(
                    id=petition_id,
                    law_firm_id=law_firm_id,
                ).first()
                if not petition:
                    return jsonify({'error': 'Petição selecionada não encontrada.'}), 404
                law_firm_document_identifier = petition.office_document_identifier
            else:
                if not law_firm_document_identifier:
                    return jsonify({'error': 'Informe o identificador do documento para o escritório.'}), 400
                if len(law_firm_document_identifier) > 96:
                    return jsonify({'error': 'O identificador do documento deve ter no máximo 96 caracteres.'}), 400
            
            # Salvar arquivo principal
            upload_dir = _create_upload_directory(law_firm_id, 'revisions')
            filename = secure_filename(main_file.filename)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S_')
            filename = timestamp + filename
            filepath = upload_dir / filename
            main_file.save(str(filepath))
            
            # Processar documentos auxiliares
            auxiliary_count = 0
            auxiliary_files = []
            if 'auxiliary_documents' in request.files:
                files = request.files.getlist('auxiliary_documents')
                for file in files:
                    if file and file.filename and allowed_file(file.filename, ALLOWED_AUXILIARY_EXTENSIONS):
                        aux_filename = secure_filename(file.filename)
                        aux_filename = timestamp + f'aux_{auxiliary_count}_' + aux_filename
                        aux_filepath = upload_dir / aux_filename
                        file.save(str(aux_filepath))
                        auxiliary_files.append({
                            'name': file.filename,
                            'path': str(aux_filepath)
                        })
                        auxiliary_count += 1

            benefits_spreadsheet = None
            if 'benefits_spreadsheet' in request.files:
                spreadsheet_file = request.files['benefits_spreadsheet']
                if spreadsheet_file and spreadsheet_file.filename:
                    if not allowed_file(spreadsheet_file.filename, ALLOWED_BENEFITS_SPREADSHEET_EXTENSIONS):
                        return jsonify({'error': 'A planilha de benefícios deve estar em formato .xlsx.'}), 400

                    spreadsheet_filename = secure_filename(spreadsheet_file.filename)
                    spreadsheet_filename = timestamp + 'benefits_' + spreadsheet_filename
                    spreadsheet_filepath = upload_dir / spreadsheet_filename
                    spreadsheet_file.save(str(spreadsheet_filepath))
                    benefits_spreadsheet = {
                        'name': spreadsheet_file.filename,
                        'path': str(spreadsheet_filepath),
                    }
            
            # Verificar se há análise comparativa
            compared_document = None
            comparative_analysis = False
            if 'compared_document' in request.files:
                compared_file = request.files['compared_document']
                if compared_file and compared_file.filename:
                    compared_extension = _get_file_extension(compared_file.filename)
                    if compared_extension == '.doc':
                        return jsonify({'error': 'Arquivos .doc não são suportados no FAP Review. Envie em PDF ou DOCX.'}), 400

                if compared_file and compared_file.filename and allowed_file(compared_file.filename, ALLOWED_DOCUMENT_EXTENSIONS):
                    compared_filename = secure_filename(compared_file.filename)
                    compared_filename = timestamp + 'compared_' + compared_filename
                    compared_filepath = upload_dir / compared_filename
                    compared_file.save(str(compared_filepath))
                    compared_document = str(compared_filepath)
                    comparative_analysis = True

            if not petition:
                petition = FapReviewPetition.query.filter_by(
                    law_firm_id=law_firm_id,
                    office_document_identifier=law_firm_document_identifier,
                ).first()

            petition_created = False
            petition_created_title = ''
            if not petition:
                petition = FapReviewPetition(
                    law_firm_id=law_firm_id,
                    created_by_id=user_id,
                    office_document_identifier=law_firm_document_identifier,
                    title=_build_petition_title(
                        petition_title,
                        fallback_filename=main_file.filename,
                        fallback_identifier=law_firm_document_identifier,
                    ),
                    workflow_status='in_review',
                )
                db.session.add(petition)
                db.session.flush()
                petition_created = True
                petition_created_title = petition.title

            next_revision_number = FapReviewExecution.query.filter_by(
                petition_id=petition.id,
                execution_type='revision',
            ).count() + 1
            
            # Criar registro de execução
            execution = FapReviewExecution(
                law_firm_id=law_firm_id,
                user_id=user_id,
                petition_id=petition.id,
                execution_type='revision',
                status='processing',
                revision_number=next_revision_number,
                main_document_path=str(filepath),
                main_document_filename=main_file.filename,
                law_firm_document_identifier=law_firm_document_identifier,
                auxiliary_documents_count=auxiliary_count,
                auxiliary_documents_json=json.dumps(auxiliary_files),
                comparative_analysis=comparative_analysis,
                compared_document_path=compared_document
            )
            db.session.add(execution)
            db.session.flush()
            petition.latest_revision_id = execution.id
            petition.revision_count = max(petition.revision_count or 0, next_revision_number)
            petition.workflow_status = 'in_review'
            petition.updated_at = datetime.utcnow()
            db.session.commit()

            if petition_created:
                _log_audit(
                    law_firm_id,
                    'petition_created',
                    'petition',
                    petition.id,
                    f'Petição criada: {petition_created_title}',
                )
            
            # Log de auditoria
            _log_audit(law_firm_id, 'revision_started', 'execution', execution.id,
                      f'Revisão iniciada: {main_file.filename}')
            
            # ===== INVOCAR AGENTE REVISOR =====
            try:
                petition_file_path = str(filepath)

                if not Path(petition_file_path).exists():
                    raise ValueError("Arquivo principal não encontrado para análise")

                compared_file_path = None
                if comparative_analysis and compared_document:
                    compared_file_path = compared_document
                    if not Path(compared_file_path).exists():
                        raise ValueError("Arquivo comparado não encontrado para análise")
                
                # Executar agente revisor
                _execute_reviewer_agent(
                    execution.id,
                    law_firm_id,
                    petition_file_path,
                    compared_file_path,
                    benefits_spreadsheet=benefits_spreadsheet,
                )
            
            except Exception as agent_error:
                current_app.logger.error(f"Erro na execução do agente: {agent_error}")
                # A função _execute_reviewer_agent já atualiza o status para 'failed'
                # Mas vamos garantir que foi marcado como falha
                execution = FapReviewExecution.query.get(execution.id)
                if execution and execution.status == 'processing':
                    execution.status = 'failed'
                    execution.error_message = str(agent_error)
                    db.session.commit()
            
            return jsonify({
                'success': True,
                'petition_id': petition.id,
                'execution_id': execution.id,
                'message': 'Revisão iniciada com sucesso'
            })
        
        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f"Erro ao processar upload: {e}")
            return jsonify({'error': str(e)}), 500
    
    # GET - Exibir página
    setting = _get_fap_setting(law_firm_id)
    petitions = FapReviewPetition.query.filter_by(
        law_firm_id=law_firm_id,
    ).order_by(
        FapReviewPetition.updated_at.desc(),
        FapReviewPetition.id.desc(),
    ).all()
    selected_petition_id = request.args.get('petition_id', type=int)
    selected_petition = None
    if selected_petition_id:
        selected_petition = FapReviewPetition.query.filter_by(
            id=selected_petition_id,
            law_firm_id=law_firm_id,
        ).first()

    return render_template(
        'fap_review/revision.html',
        setting=setting,
        petitions=petitions,
        selected_petition=selected_petition,
        petition_status_labels=PETITION_WORKFLOW_STATUSES,
    )


@fap_review_bp.route('/revision/<int:execution_id>', methods=['GET'])
@require_law_firm
def revision_result(execution_id: int):
    """Exibe resultado de uma revisão"""
    law_firm_id = get_current_law_firm_id()
    
    execution = FapReviewExecution.query.filter_by(
        id=execution_id,
        law_firm_id=law_firm_id
    ).first_or_404()
    petition = execution.petition
    document_identifier = (
        petition.office_document_identifier
        if petition and petition.office_document_identifier
        else execution.law_firm_document_identifier
    )
    
    result_data = {}
    if execution.result_json:
        try:
            result_data = json.loads(execution.result_json)
        except json.JSONDecodeError:
            result_data = {}

    ignored_finding_indices: list[int] = []
    if document_identifier and result_data.get('findings'):
        ignored_fingerprints = _get_ignored_finding_fingerprints(
            law_firm_id=law_firm_id,
            law_firm_document_identifier=document_identifier,
        )
        ignored_finding_indices = [
            index
            for index, finding in enumerate(result_data.get('findings') or [], start=1)
            if _build_finding_fingerprint(finding) in ignored_fingerprints
        ]

    prior_revision_count = 0
    if execution.petition_id:
        prior_revision_count = FapReviewExecution.query.filter(
            FapReviewExecution.law_firm_id == law_firm_id,
            FapReviewExecution.execution_type == 'revision',
            FapReviewExecution.petition_id == execution.petition_id,
            FapReviewExecution.id != execution.id,
        ).count()
    elif execution.law_firm_document_identifier:
        prior_revision_count = FapReviewExecution.query.filter(
            FapReviewExecution.law_firm_id == law_firm_id,
            FapReviewExecution.execution_type == 'revision',
            FapReviewExecution.law_firm_document_identifier == execution.law_firm_document_identifier,
            FapReviewExecution.id != execution.id,
        ).count()

    manual_reference = _get_active_reference(law_firm_id, 'manual_fap')
    manual_content = (manual_reference.content if manual_reference else '').strip()
    
    return render_template('fap_review/revision_result.html',
                          execution=execution,
                          petition=petition,
                          petition_status_badge=_build_petition_status_badge(petition.workflow_status) if petition else None,
                          result_data=result_data,
                          prior_revision_count=prior_revision_count,
                          ignored_finding_indices=ignored_finding_indices,
                          manual_content=manual_content,
                          manual_reference=manual_reference)


@fap_review_bp.route('/petitions/<int:petition_id>', methods=['GET'])
@require_law_firm
def petition_detail(petition_id: int):
    """Exibe o histórico agregado de revisões de uma petição."""
    law_firm_id = get_current_law_firm_id()

    petition = FapReviewPetition.query.filter_by(
        id=petition_id,
        law_firm_id=law_firm_id,
    ).first_or_404()

    revisions = FapReviewExecution.query.filter_by(
        law_firm_id=law_firm_id,
        petition_id=petition.id,
        execution_type='revision',
    ).order_by(
        FapReviewExecution.created_at.desc(),
        FapReviewExecution.id.desc(),
    ).all()

    return render_template(
        'fap_review/petition_detail.html',
        petition=petition,
        revisions=revisions,
        status_badge=_build_petition_status_badge(petition.workflow_status),
        petition_status_labels=PETITION_WORKFLOW_STATUSES,
    )


@fap_review_bp.route('/petitions/<int:petition_id>/status', methods=['POST'])
@require_law_firm
def petition_update_status(petition_id: int):
    """Permite controle manual do status agregado da petição."""
    law_firm_id = get_current_law_firm_id()

    petition = FapReviewPetition.query.filter_by(
        id=petition_id,
        law_firm_id=law_firm_id,
    ).first_or_404()

    payload = request.get_json(silent=True) or {}
    new_status = str(payload.get('workflow_status') or '').strip()
    if new_status not in PETITION_WORKFLOW_STATUSES:
        return jsonify({'error': 'Status da petição inválido.'}), 400

    old_status = petition.workflow_status
    if new_status == old_status:
        return jsonify({'success': True, 'workflow_status': new_status})

    try:
        petition.workflow_status = new_status
        petition.updated_at = datetime.utcnow()
        db.session.commit()
    except Exception as error:
        db.session.rollback()
        current_app.logger.error(f'Erro ao atualizar status da petição: {error}')
        return jsonify({'error': 'Não foi possível atualizar o status da petição.'}), 500

    _log_audit(
        law_firm_id,
        'petition_status_updated',
        'petition',
        petition.id,
        f'Status alterado de {PETITION_WORKFLOW_STATUSES.get(old_status, old_status)} para {PETITION_WORKFLOW_STATUSES.get(new_status, new_status)}',
        old_value=old_status,
        new_value=new_status,
    )

    return jsonify({'success': True, 'workflow_status': new_status})


@fap_review_bp.route('/revision/<int:execution_id>/findings/<int:finding_index>/dismiss-feedback', methods=['POST'])
@require_law_firm
def revision_finding_feedback(execution_id: int, finding_index: int):
    """Persiste feedback de ponto marcado como não útil para futuras revisões do mesmo documento."""
    law_firm_id = get_current_law_firm_id()
    user_id = session.get('user_id')

    execution = FapReviewExecution.query.filter_by(
        id=execution_id,
        law_firm_id=law_firm_id,
    ).first_or_404()

    if not execution.law_firm_document_identifier:
        return jsonify({'error': 'Esta revisão não possui identificador de documento para persistir feedback.'}), 400

    try:
        payload = json.loads(execution.result_json or '{}')
    except json.JSONDecodeError:
        return jsonify({'error': 'Resultado da revisão inválido para registrar feedback.'}), 400

    findings = payload.get('findings') or []
    if finding_index < 1 or finding_index > len(findings):
        return jsonify({'error': 'Ponto de atenção não encontrado nesta revisão.'}), 404

    finding = findings[finding_index - 1]
    finding_fingerprint = _build_finding_fingerprint(finding)
    if not finding_fingerprint:
        return jsonify({'error': 'Não foi possível identificar o ponto de atenção selecionado.'}), 400

    request_payload = request.get_json(silent=True) or {}
    dismissed = bool(request_payload.get('dismissed'))

    existing_feedback = FapReviewIgnoredFinding.query.filter_by(
        law_firm_id=law_firm_id,
        law_firm_document_identifier=execution.law_firm_document_identifier,
        finding_fingerprint=finding_fingerprint,
    ).first()

    description = str(finding.get('description') or '').strip()
    audit_action = None
    audit_description = None

    try:
        if dismissed:
            if not existing_feedback:
                db.session.add(FapReviewIgnoredFinding(
                    law_firm_id=law_firm_id,
                    law_firm_document_identifier=execution.law_firm_document_identifier,
                    source_execution_id=execution.id,
                    created_by_id=user_id,
                    finding_fingerprint=finding_fingerprint,
                    finding_description=description,
                ))
            audit_action = 'revision_finding_dismissed'
            audit_description = f'Ponto marcado como não útil: {description[:200]}'
        elif existing_feedback:
            db.session.delete(existing_feedback)
            audit_action = 'revision_finding_dismiss_undone'
            audit_description = f'Ponto voltou a ser considerado no histórico: {description[:200]}'

        db.session.commit()
    except Exception as error:
        db.session.rollback()
        current_app.logger.error(f'Erro ao persistir feedback do ponto de atenção: {error}')
        return jsonify({'error': 'Não foi possível salvar o feedback do ponto de atenção.'}), 500

    if audit_action and audit_description:
        try:
            _log_audit(
                law_firm_id,
                audit_action,
                'execution',
                execution.id,
                audit_description,
            )
        except Exception as audit_error:
            current_app.logger.warning(f'Feedback salvo, mas falha ao registrar auditoria: {audit_error}')

    return jsonify({'success': True, 'dismissed': dismissed})


@fap_review_bp.route('/revision/<int:execution_id>/document/main', methods=['GET'])
@require_law_firm
def revision_main_document(execution_id: int):
    """Abre o documento principal salvo da execução de revisão."""
    law_firm_id = get_current_law_firm_id()

    execution = FapReviewExecution.query.filter_by(
        id=execution_id,
        law_firm_id=law_firm_id
    ).first_or_404()

    file_path = str(execution.main_document_path or '').strip()
    if not file_path:
        flash('Documento principal não disponível para esta execução.', 'warning')
        return redirect(url_for('fap_review.revision_result', execution_id=execution_id))

    path = Path(file_path)
    if not path.exists() or not path.is_file():
        flash('Arquivo da revisão não foi encontrado no armazenamento.', 'error')
        return redirect(url_for('fap_review.revision_result', execution_id=execution_id))

    return send_file(
        path,
        as_attachment=False,
        download_name=execution.main_document_filename or path.name,
    )


@fap_review_bp.route('/training', methods=['GET', 'POST'])
@require_law_firm
@require_admin_user
def training():
    """Página de gerenciamento de treinamento e evolução"""
    law_firm_id = get_current_law_firm_id()
    user_id = session.get('user_id')
    setting = _get_fap_setting(law_firm_id)

    training_preview = None
    preview_execution_id = None
    preview_files = {}
    apply_result = None
    
    if request.method == 'POST':
        action = request.form.get('action', '').strip().lower()

        if action == 'compare':
            try:
                if 'original_document' not in request.files or 'revised_document' not in request.files:
                    flash('Envie os dois arquivos para comparação.', 'warning')
                    return redirect(url_for('fap_review.training'))

                original_file = request.files['original_document']
                revised_file = request.files['revised_document']

                if not original_file.filename or not revised_file.filename:
                    flash('Selecione ambos os arquivos (original e revisado).', 'warning')
                    return redirect(url_for('fap_review.training'))

                if _get_file_extension(original_file.filename) == '.doc' or _get_file_extension(revised_file.filename) == '.doc':
                    flash('Arquivos .doc não são suportados no FAP Review. Envie em PDF ou DOCX.', 'error')
                    return redirect(url_for('fap_review.training'))

                if not allowed_file(original_file.filename, ALLOWED_DOCUMENT_EXTENSIONS):
                    flash('Arquivo original com extensão não permitida.', 'error')
                    return redirect(url_for('fap_review.training'))

                if not allowed_file(revised_file.filename, ALLOWED_DOCUMENT_EXTENSIONS):
                    flash('Arquivo revisado com extensão não permitida.', 'error')
                    return redirect(url_for('fap_review.training'))

                upload_dir = _create_upload_directory(law_firm_id, 'training')
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S_')

                original_filename_safe = timestamp + 'original_' + secure_filename(original_file.filename)
                revised_filename_safe = timestamp + 'revised_' + secure_filename(revised_file.filename)

                original_path = upload_dir / original_filename_safe
                revised_path = upload_dir / revised_filename_safe

                original_file.save(str(original_path))
                revised_file.save(str(revised_path))

                original_text = _extract_text_from_document(str(original_path))
                revised_text = _extract_text_from_document(str(revised_path))

                training_identity_prompt = FapReviewPromptVersion.query.filter_by(
                    law_firm_id=law_firm_id,
                    prompt_type='training_identity',
                    is_active=True,
                ).first()
                training_rules_prompt = FapReviewPromptVersion.query.filter_by(
                    law_firm_id=law_firm_id,
                    prompt_type='training_rules',
                    is_active=True,
                ).first()
                training_prompt = FapReviewPromptVersion.query.filter_by(
                    law_firm_id=law_firm_id,
                    prompt_type='training_prompt',
                    is_active=True,
                ).first()

                subagent = FapTrainingApplySubAgent(
                    openai_api_key=os.getenv('OPENAI_API_KEY'),
                    model=setting.training_model,
                    temperature=min(max(setting.training_temperature, 0.0), 1.0),
                )

                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                try:
                    extract_result = loop.run_until_complete(
                        subagent.build_comparison_extract(
                            original_text=original_text,
                            revised_text=revised_text,
                            training_identity=training_identity_prompt.content if training_identity_prompt else '',
                            training_rules=training_rules_prompt.content if training_rules_prompt else '',
                            training_prompt=training_prompt.content if training_prompt else '',
                        )
                    )
                finally:
                    loop.close()

                preview_payload = {
                    'stage': 'preview',
                    'extract': extract_result.model_dump(),
                    'source_files': {
                        'original_filename': original_file.filename,
                        'revised_filename': revised_file.filename,
                        'original_path': str(original_path),
                        'revised_path': str(revised_path),
                    },
                }

                execution = FapReviewExecution(
                    law_firm_id=law_firm_id,
                    user_id=user_id,
                    execution_type='training',
                    status='pending',
                    main_document_path=str(original_path),
                    main_document_filename=original_file.filename,
                    comparative_analysis=True,
                    compared_document_path=str(revised_path),
                    result_json=json.dumps(preview_payload, ensure_ascii=False),
                )
                db.session.add(execution)
                db.session.commit()

                _log_audit(
                    law_firm_id,
                    'training_preview_generated',
                    'execution',
                    execution.id,
                    'Extrato de comparação gerado para confirmação humana',
                )

                training_preview = extract_result.model_dump()
                preview_execution_id = execution.id
                preview_files = {
                    'original_filename': original_file.filename,
                    'revised_filename': revised_file.filename,
                }

                flash('Extrato da comparação gerado. Revise e confirme para treinar.', 'success')

            except Exception as e:
                db.session.rollback()
                current_app.logger.error(f"Erro ao gerar extrato de treinamento: {e}")
                flash(f'Erro ao comparar documentos: {str(e)}', 'error')

        elif action == 'apply':
            preview_id = request.form.get('preview_execution_id', type=int)
            if not preview_id:
                flash('Execução de prévia não informada.', 'warning')
                return redirect(url_for('fap_review.training'))

            execution = FapReviewExecution.query.filter_by(
                id=preview_id,
                law_firm_id=law_firm_id,
                execution_type='training',
            ).first()

            if not execution:
                flash('Prévia de treinamento não encontrada.', 'error')
                return redirect(url_for('fap_review.training'))

            try:
                payload = json.loads(execution.result_json or '{}')
                extract_data = payload.get('extract') or {}
                source_files = payload.get('source_files') or {}

                if not extract_data:
                    flash('A prévia não possui extrato para aplicação.', 'warning')
                    return redirect(url_for('fap_review.training'))

                training_update_policy_prompt = FapReviewPromptVersion.query.filter_by(
                    law_firm_id=law_firm_id,
                    prompt_type='training_update_policy',
                    is_active=True,
                ).first()

                manual_ref = _get_active_reference(law_firm_id, 'manual_fap')
                cases_ref = _get_active_reference(law_firm_id, 'casos_referencia')

                manual_content = (manual_ref.content if manual_ref else '').strip()
                cases_content = (cases_ref.content if cases_ref else '').strip()

                subagent = FapTrainingApplySubAgent(
                    openai_api_key=os.getenv('OPENAI_API_KEY'),
                    model=setting.training_model,
                    temperature=min(max(setting.training_temperature, 0.0), 1.0),
                )

                trainer = FapTrainingEvolutionAgent(
                    openai_api_key=os.getenv('OPENAI_API_KEY'),
                    model=setting.training_model,
                    temperature=min(max(setting.training_temperature, 0.0), 1.0),
                )

                manual_version = '1.0.0'
                if manual_ref and manual_ref.version_number:
                    manual_version = f'1.0.{manual_ref.version_number}'

                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                try:
                    apply_payload = loop.run_until_complete(
                        subagent.build_apply_payload(
                            extract_data=extract_data,
                            training_update_policy=(
                                training_update_policy_prompt.content
                                if training_update_policy_prompt
                                else ''
                            ),
                            manual_version=manual_version,
                            original_filename=source_files.get('original_filename', ''),
                            revised_filename=source_files.get('revised_filename', ''),
                        )
                    )
                finally:
                    loop.close()

                timestamp_label = now_sp().strftime('%d/%m/%Y %H:%M')

                should_update_manual = bool(apply_payload.should_update_manual and setting.auto_update_manual)
                should_update_cases = bool(apply_payload.should_update_cases and setting.auto_update_cases)

                manual_patch = (apply_payload.manual_patch_markdown or '').strip()
                if manual_patch and should_update_manual:
                    new_manual_content = (
                        manual_content + '\n\n' + manual_patch
                        if manual_content else manual_patch
                    )
                    manual_new = _append_reference_version(
                        law_firm_id=law_firm_id,
                        user_id=user_id,
                        reference_type='manual_fap',
                        new_content=new_manual_content,
                        activate=True,
                    )
                    manual_updated = True
                else:
                    manual_new = None
                    manual_updated = False

                cases_patch = (apply_payload.case_reference_markdown or '').strip()
                if cases_patch and should_update_cases:
                    new_cases_content = (
                        cases_content + '\n\n' + cases_patch
                        if cases_content else cases_patch
                    )
                    cases_new = _append_reference_version(
                        law_firm_id=law_firm_id,
                        user_id=user_id,
                        reference_type='casos_referencia',
                        new_content=new_cases_content,
                        activate=True,
                    )
                    cases_updated = True
                else:
                    cases_new = None
                    cases_updated = False

                semantic_version_new = trainer._increment_version(
                    manual_version,
                    apply_payload.version_increment or 'patch',
                )

                execution.status = 'completed'
                execution.completed_at = datetime.utcnow()
                execution.updated_at = datetime.utcnow()
                execution.result_json = json.dumps(
                    {
                        'stage': 'applied',
                        'applied_at': timestamp_label,
                        'extract': extract_data,
                        'source_files': source_files,
                        'training_result': {
                            'manual_updates_generated': manual_updated,
                            'case_reference_generated': cases_updated,
                            'manual_version_new': semantic_version_new,
                            'approval_required': setting.require_approval_before_publish,
                            'message': apply_payload.message,
                            'reference_versions': {
                                'manual_fap': manual_new.version_number if manual_new else None,
                                'casos_referencia': cases_new.version_number if cases_new else None,
                            },
                        },
                    },
                    ensure_ascii=False,
                )

                db.session.commit()

                _log_audit(
                    law_firm_id,
                    'training_applied',
                    'execution',
                    execution.id,
                    'Treinamento aplicado com confirmação humana',
                )

                apply_result = {
                    'manual_updated': manual_updated,
                    'cases_updated': cases_updated,
                    'manual_version_new': semantic_version_new,
                    'message': apply_payload.message,
                }
                flash('Treinamento aplicado com sucesso.', 'success')

            except Exception as e:
                db.session.rollback()
                current_app.logger.error(f"Erro ao aplicar treinamento: {e}")

                execution.status = 'failed'
                execution.error_message = str(e)
                execution.completed_at = datetime.utcnow()
                db.session.commit()

                flash(f'Erro ao aplicar treinamento: {str(e)}', 'error')
    
    # Últimas execuções de treinamento
    recent_training = FapReviewExecution.query.filter_by(
        law_firm_id=law_firm_id,
        execution_type='training'
    ).order_by(FapReviewExecution.created_at.desc()).limit(10).all()
    
    return render_template('fap_review/training.html',
                          setting=setting,
                          recent_training=recent_training,
                          training_preview=training_preview,
                          preview_execution_id=preview_execution_id,
                          preview_files=preview_files,
                          apply_result=apply_result)


@fap_review_bp.route('/settings', methods=['GET', 'POST'])
@require_law_firm
@require_admin_user
def settings():
    """Página de configurações do módulo"""
    law_firm_id = get_current_law_firm_id()
    user_id = session.get('user_id')
    
    setting = _get_fap_setting(law_firm_id)
    
    if request.method == 'POST':
        try:
            data = request.json
            
            # Atualizar configurações
            setting.reviewer_model = data.get('reviewer_model', setting.reviewer_model)
            setting.training_model = data.get('training_model', setting.training_model)
            # Normalizar temperatura (converter vírgula para ponto se necessário)
            reviewer_temp_str = str(data.get('reviewer_temperature', setting.reviewer_temperature)).replace(',', '.')
            setting.reviewer_temperature = float(reviewer_temp_str)
            training_temp_str = str(data.get('training_temperature', setting.training_temperature)).replace(',', '.')
            setting.training_temperature = float(training_temp_str)
            setting.auto_update_manual = data.get('auto_update_manual', setting.auto_update_manual)
            setting.auto_update_cases = data.get('auto_update_cases', setting.auto_update_cases)
            setting.require_approval_before_publish = data.get('require_approval_before_publish', setting.require_approval_before_publish)
            setting.enable_continuous_learning = data.get('enable_continuous_learning', setting.enable_continuous_learning)
            setting.reviewer_enabled = data.get('reviewer_enabled', setting.reviewer_enabled)
            setting.training_enabled = data.get('training_enabled', setting.training_enabled)
            setting.updated_at = now_sp()
            
            db.session.commit()
            
            # Log de auditoria
            _log_audit(law_firm_id, 'settings_updated', 'setting', setting.id,
                      'Configurações do FAP Review atualizadas')
            
            return jsonify({'success': True, 'message': 'Configurações salvas com sucesso'})
        
        except Exception as e:
            db.session.rollback()
            return jsonify({'error': str(e)}), 500
    
    # GET - Exibir página
    available_models, model_options_error = fetch_openrouter_text_models_for_info(
        selected_model=setting.reviewer_model or '',
        fallback_model=setting.reviewer_model or '',
    )
    return render_template('fap_review/settings.html',
                           setting=setting,
                           available_models=available_models,
                           model_options_error=model_options_error)


@fap_review_bp.route('/settings/prompts/type/<string:prompt_type>/edit', methods=['GET'])
@require_law_firm
@require_admin_user
def edit_prompt_by_type(prompt_type: str):
    """Abre edição do prompt pela tipagem, criando versão inicial se necessário."""
    law_firm_id = get_current_law_firm_id()

    valid_types = {
        'revisor_identity',
        'revisor_rules',
        'revisor_output_format',
        'training_identity',
        'training_rules',
        'training_prompt',
        'training_update_policy',
    }

    if prompt_type not in valid_types:
        flash('Tipo de prompt inválido', 'error')
        return redirect(url_for('fap_review.settings'))

    prompt = FapReviewPromptVersion.query.filter_by(
        law_firm_id=law_firm_id,
        prompt_type=prompt_type,
        is_active=True,
    ).order_by(FapReviewPromptVersion.version_number.desc()).first()

    if not prompt:
        prompt = FapReviewPromptVersion.query.filter_by(
        law_firm_id=law_firm_id,
        prompt_type=prompt_type,
    ).order_by(FapReviewPromptVersion.version_number.desc()).first()

    if not prompt:
        prompt = FapReviewPromptVersion(
            law_firm_id=law_firm_id,
            version_number=1,
            prompt_type=prompt_type,
            content='',
            is_active=True,
            created_by_id=session.get('user_id'),
        )
        db.session.add(prompt)
        db.session.commit()

        _log_audit(
            law_firm_id,
            'prompt_created',
            'prompt',
            prompt.id,
            f'Prompt inicial criado para {prompt_type}'
        )

    return redirect(url_for('fap_review.edit_prompt', prompt_version_id=prompt.id))


@fap_review_bp.route('/settings/references/type/<string:reference_type>/edit', methods=['GET'])
@require_law_firm
@require_admin_user
def edit_reference_by_type(reference_type: str):
    """Abre edição da referência pela tipagem, criando versão inicial se necessário."""
    law_firm_id = get_current_law_firm_id()

    valid_types = {
        'manual_fap',
        'casos_referencia',
        'project_instructions',
    }

    if reference_type not in valid_types:
        flash('Tipo de referência inválido', 'error')
        return redirect(url_for('fap_review.settings'))

    reference = FapReviewReferenceVersion.query.filter_by(
        law_firm_id=law_firm_id,
        reference_type=reference_type,
        is_active=True,
    ).order_by(FapReviewReferenceVersion.version_number.desc()).first()

    if not reference:
        reference = FapReviewReferenceVersion.query.filter_by(
        law_firm_id=law_firm_id,
        reference_type=reference_type,
    ).order_by(FapReviewReferenceVersion.version_number.desc()).first()

    if not reference:
        reference = FapReviewReferenceVersion(
            law_firm_id=law_firm_id,
            version_number=1,
            reference_type=reference_type,
            content='',
            is_active=True,
            created_by_id=session.get('user_id'),
        )
        db.session.add(reference)
        db.session.commit()

        _log_audit(
            law_firm_id,
            'reference_created',
            'reference',
            reference.id,
            f'Referência inicial criada para {reference_type}'
        )

    return redirect(url_for('fap_review.edit_reference', reference_version_id=reference.id))


# ═══════════════════════════════════════════════════════════════════════════════
# ROTAS DE PROMPTS E REFERENCIAS
# ═══════════════════════════════════════════════════════════════════════════════


@fap_review_bp.route('/settings/prompts', methods=['GET'])
@require_law_firm
@require_admin_user
def list_prompts():
    """Lista prompts por tipo"""
    law_firm_id = get_current_law_firm_id()
    
    prompt_types = [
        'revisor_identity',
        'revisor_rules',
        'revisor_output_format',
        'training_identity',
        'training_rules',
        'training_prompt',
        'training_update_policy'
    ]
    
    prompts = {}
    for ptype in prompt_types:
        versions = FapReviewPromptVersion.query.filter_by(
            law_firm_id=law_firm_id,
            prompt_type=ptype
        ).order_by(FapReviewPromptVersion.version_number.desc()).all()
        prompts[ptype] = versions
    
    return jsonify({
        'prompts': {
            k: [{'version': v.version_number, 'is_active': v.is_active, 'created_at': v.created_at.isoformat()} for v in v]
            for k, v in prompts.items()
        }
    })


@fap_review_bp.route('/settings/prompts/<int:prompt_version_id>', methods=['GET', 'POST'])
@require_law_firm
@require_admin_user
def edit_prompt(prompt_version_id: int):
    """Edita um prompt específico"""
    law_firm_id = get_current_law_firm_id()
    
    prompt = FapReviewPromptVersion.query.filter_by(
        id=prompt_version_id,
        law_firm_id=law_firm_id
    ).first_or_404()

    # Evita abrir versão inativa por padrão; direciona para a ativa mais recente.
    if request.method == 'GET' and not prompt.is_active:
        active_prompt = FapReviewPromptVersion.query.filter_by(
            law_firm_id=law_firm_id,
            prompt_type=prompt.prompt_type,
            is_active=True,
        ).order_by(FapReviewPromptVersion.version_number.desc()).first()
        if active_prompt and active_prompt.id != prompt.id:
            return redirect(url_for('fap_review.edit_prompt', prompt_version_id=active_prompt.id))
    
    # Carregar todas as versões deste tipo
    all_versions = FapReviewPromptVersion.query.filter_by(
        law_firm_id=law_firm_id,
        prompt_type=prompt.prompt_type
    ).order_by(FapReviewPromptVersion.version_number.desc()).all()

    is_read_only_prompt = prompt.prompt_type in READ_ONLY_PROMPT_TYPES
    
    if request.method == 'POST':
        if is_read_only_prompt:
            return jsonify({'error': 'Este prompt é somente leitura e não pode ser editado.'}), 403

        try:
            content = request.json.get('content', '')
            
            # Criar nova versão
            new_version = FapReviewPromptVersion(
                law_firm_id=law_firm_id,
                version_number=prompt.version_number + 1,
                prompt_type=prompt.prompt_type,
                content=content,
                is_active=False,
                created_by_id=session.get('user_id')
            )
            db.session.add(new_version)
            db.session.commit()
            
            # Log
            _log_audit(law_firm_id, 'prompt_updated', 'prompt', new_version.id,
                      f'Nova versão criada: v{new_version.version_number}')
            
            return jsonify({'success': True, 'version_id': new_version.id})
        
        except Exception as e:
            db.session.rollback()
            return jsonify({'error': str(e)}), 500
    
    # GET
    return render_template(
        'fap_review/edit_prompt.html',
        prompt=prompt,
        versions=all_versions,
        is_read_only_prompt=is_read_only_prompt,
    )


@fap_review_bp.route('/settings/prompts/<int:prompt_version_id>/activate', methods=['POST'])
@require_law_firm
@require_admin_user
def activate_prompt(prompt_version_id: int):
    """Ativa uma versão de prompt"""
    law_firm_id = get_current_law_firm_id()
    
    prompt = FapReviewPromptVersion.query.filter_by(
        id=prompt_version_id,
        law_firm_id=law_firm_id
    ).first_or_404()

    if prompt.prompt_type in READ_ONLY_PROMPT_TYPES:
        return jsonify({'error': 'Este prompt é somente leitura e não permite ativação manual.'}), 403
    
    try:
        # Desativar outras versões do mesmo tipo
        FapReviewPromptVersion.query.filter_by(
            law_firm_id=law_firm_id,
            prompt_type=prompt.prompt_type,
            is_active=True
        ).update({'is_active': False})
        
        # Ativar esta versão
        prompt.is_active = True
        db.session.commit()
        
        _log_audit(law_firm_id, 'prompt_activated', 'prompt', prompt.id,
                  f'Versão v{prompt.version_number} ativada')
        
        return jsonify({'success': True})
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@fap_review_bp.route('/settings/references', methods=['GET'])
@require_law_firm
@require_admin_user
def list_references():
    """Lista documentos de referência"""
    law_firm_id = get_current_law_firm_id()
    
    reference_types = ['manual_fap', 'casos_referencia', 'project_instructions']
    
    references = {}
    for rtype in reference_types:
        versions = FapReviewReferenceVersion.query.filter_by(
            law_firm_id=law_firm_id,
            reference_type=rtype
        ).order_by(FapReviewReferenceVersion.version_number.desc()).all()
        references[rtype] = versions
    
    return jsonify({
        'references': {
            k: [{'version': v.version_number, 'is_active': v.is_active, 'created_at': v.created_at.isoformat()} for v in v]
            for k, v in references.items()
        }
    })


@fap_review_bp.route('/settings/references/<int:reference_version_id>', methods=['GET', 'POST'])
@require_law_firm
@require_admin_user
def edit_reference(reference_version_id: int):
    """Edita um documento de referência"""
    law_firm_id = get_current_law_firm_id()
    
    reference = FapReviewReferenceVersion.query.filter_by(
        id=reference_version_id,
        law_firm_id=law_firm_id
    ).first_or_404()

    # Evita abrir versão inativa por padrão; direciona para a ativa mais recente.
    if request.method == 'GET' and not reference.is_active:
        active_reference = FapReviewReferenceVersion.query.filter_by(
            law_firm_id=law_firm_id,
            reference_type=reference.reference_type,
            is_active=True,
        ).order_by(FapReviewReferenceVersion.version_number.desc()).first()
        if active_reference and active_reference.id != reference.id:
            return redirect(url_for('fap_review.edit_reference', reference_version_id=active_reference.id))
    
    # Carregar todas as versões deste tipo
    all_versions = FapReviewReferenceVersion.query.filter_by(
        law_firm_id=law_firm_id,
        reference_type=reference.reference_type
    ).order_by(FapReviewReferenceVersion.version_number.desc()).all()

    is_read_only_reference = reference.reference_type in READ_ONLY_REFERENCE_TYPES
    
    if request.method == 'POST':
        if is_read_only_reference:
            return jsonify({'error': 'Esta referência é somente leitura e não pode ser editada.'}), 403

        try:
            content = request.json.get('content', '')
            
            # Criar nova versão
            new_version = FapReviewReferenceVersion(
                law_firm_id=law_firm_id,
                version_number=reference.version_number + 1,
                reference_type=reference.reference_type,
                content=content,
                is_active=False,
                created_by_id=session.get('user_id')
            )
            db.session.add(new_version)
            db.session.commit()
            
            _log_audit(law_firm_id, 'reference_updated', 'reference', new_version.id,
                      f'Nova versão criada: v{new_version.version_number}')
            
            return jsonify({'success': True, 'version_id': new_version.id})
        
        except Exception as e:
            db.session.rollback()
            return jsonify({'error': str(e)}), 500
    
    # GET
    return render_template(
        'fap_review/edit_reference.html',
        reference=reference,
        versions=all_versions,
        is_read_only_reference=is_read_only_reference,
    )


@fap_review_bp.route('/settings/references/<int:reference_version_id>/activate', methods=['POST'])
@require_law_firm
@require_admin_user
def activate_reference(reference_version_id: int):
    """Ativa uma versão de referência"""
    law_firm_id = get_current_law_firm_id()
    
    reference = FapReviewReferenceVersion.query.filter_by(
        id=reference_version_id,
        law_firm_id=law_firm_id
    ).first_or_404()

    if reference.reference_type in READ_ONLY_REFERENCE_TYPES:
        return jsonify({'error': 'Esta referência é somente leitura e não permite ativação manual.'}), 403
    
    try:
        # Desativar outras versões do mesmo tipo
        FapReviewReferenceVersion.query.filter_by(
            law_firm_id=law_firm_id,
            reference_type=reference.reference_type,
            is_active=True
        ).update({'is_active': False})
        
        # Ativar esta versão
        reference.is_active = True
        db.session.commit()
        
        _log_audit(law_firm_id, 'reference_activated', 'reference', reference.id,
                  f'Versão v{reference.version_number} ativada')
        
        return jsonify({'success': True})
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


# ═══════════════════════════════════════════════════════════════════════════════
# ROTAS DE AUDITORIA
# ═══════════════════════════════════════════════════════════════════════════════


@fap_review_bp.route('/audit-logs', methods=['GET'])
@require_law_firm
@require_admin_user
def audit_logs():
    """Exibe logs de auditoria"""
    law_firm_id = get_current_law_firm_id()
    
    page = request.args.get('page', 1, type=int)
    logs = FapReviewAuditLog.query.filter_by(
        law_firm_id=law_firm_id
    ).order_by(FapReviewAuditLog.created_at.desc()).paginate(page=page, per_page=50)
    
    return render_template('fap_review/audit_logs.html', logs=logs)


@fap_review_bp.route('/api/audit-logs', methods=['GET'])
@require_law_firm
@require_admin_user
def api_audit_logs():
    """API para obter logs de auditoria"""
    law_firm_id = get_current_law_firm_id()
    
    action_filter = request.args.get('action', '')
    entity_filter = request.args.get('entity_type', '')
    limit = request.args.get('limit', 50, type=int)
    
    query = FapReviewAuditLog.query.filter_by(law_firm_id=law_firm_id)
    
    if action_filter:
        query = query.filter_by(action=action_filter)
    if entity_filter:
        query = query.filter_by(entity_type=entity_filter)
    
    logs = query.order_by(FapReviewAuditLog.created_at.desc()).limit(limit).all()
    
    return jsonify({
        'logs': [{
            'id': log.id,
            'action': log.action,
            'entity_type': log.entity_type,
            'user': log.user.name if log.user else 'System',
            'created_at': log.created_at.isoformat(),
            'description': log.change_description
        } for log in logs]
    })
