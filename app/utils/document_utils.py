"""
Utilitários para processamento de documentos
"""
from docx import Document
import os


def extract_text_from_docx(file_path):
    """
    Extrai texto completo de um documento DOCX
    
    Args:
        file_path: Caminho absoluto para o arquivo DOCX
        
    Returns:
        str: Texto extraído do documento
    """
    try:
        document = Document(file_path)
        text_parts = []
        
        # Extrair texto dos parágrafos
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extrair texto das tabelas
        for table in document.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = ' '.join([p.text for p in cell.paragraphs if p.text.strip()])
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append(' | '.join(row_text))
        
        return '\n\n'.join(text_parts)
    except Exception as e:
        raise Exception(f"Erro ao extrair texto do DOCX: {str(e)}")


def get_file_extension(file_path):
    """
    Retorna a extensão do arquivo em minúsculas
    
    Args:
        file_path: Caminho do arquivo
        
    Returns:
        str: Extensão sem o ponto (ex: 'docx', 'pdf')
    """
    return os.path.splitext(file_path)[1].lower().replace('.', '')


def is_docx_file(file_path):
    """
    Verifica se o arquivo é DOCX
    
    Args:
        file_path: Caminho do arquivo
        
    Returns:
        bool: True se for DOCX
    """
    extension = get_file_extension(file_path)
    return extension in ['docx', 'doc']


def render_docx_preview_html(file_path):
    """
    Converte um DOCX em HTML semântico para visualização no navegador.

    Usa mammoth, que gera apenas tags próprias (títulos, parágrafos, negrito,
    tabelas) escapando o texto do documento — seguro para renderizar com |safe.

    Args:
        file_path: Caminho absoluto para o arquivo DOCX

    Returns:
        str: HTML do conteúdo do documento

    Raises:
        ValueError: se o arquivo não puder ser convertido
    """
    import mammoth

    try:
        with open(file_path, 'rb') as docx_file:
            return mammoth.convert_to_html(docx_file).value
    except Exception as e:
        raise ValueError(f"Erro ao converter DOCX para HTML: {str(e)}")


def strip_html_text(value):
    """Texto puro a partir de HTML: remove blocos <style>/<script>, tags e
    entidades, normalizando espaços.

    O teor de alguns tribunais (ex.: TJSC) vem em HTML completo; um striptags
    simples preserva o conteúdo interno dos blocos de estilo como texto.
    Usado pelas telas (filtro Jinja ``sem_estilos`` + striptags) e pelo MCP.
    """
    import re
    from markupsafe import Markup

    if not value:
        return ''
    text = re.sub(r'(?is)<(style|script)[^>]*>.*?</\1>', ' ', str(value))
    return Markup(text).striptags()
