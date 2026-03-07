"""
Processador de recursos judiciais (pendentes).

Uso:
  python scripts/process_judicial_appeals.py

Este script processa recursos pendentes e os gera usando IA.
"""

import os
import sys
import json
from datetime import datetime
from rich import print
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import shutil

PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
if PROJECT_ROOT not in sys.path:
    sys.path.insert(0, PROJECT_ROOT)

from main import app
from app.models import db, JudicialAppeal, JudicialSentenceAnalysis
from app.agents.legal_drafting.agent_appeal_generator import AgentAppealGenerator


def create_docx_from_appeal(appeal_content: dict, output_path: str) -> bool:
    """
    Cria documento DOCX usando modelo_documento.docx como base e inserindo dados da IA.
    
    Args:
        appeal_content: Dicionário com o conteúdo estruturado do recurso
        output_path: Caminho onde o arquivo DOCX será salvo
        
    Returns:
        bool: True se criado com sucesso, False caso contrário
    """
    try:
        # Caminho do template
        template_path = os.path.join(PROJECT_ROOT, 'templates_padrao', 'modelo_documento.docx')
        
        if not os.path.exists(template_path):
            print(f"⚠ Template não encontrado: {template_path}")
            print("  Criando documento sem template...")
            return _create_docx_manual(appeal_content, output_path)
        
        # Copiar o template para o destino (mantém formatação, cabeçalho, rodapé, etc.)
        shutil.copy2(template_path, output_path)
        print(f"✓ Template copiado: {template_path} → {output_path}")
        
        # Abrir o documento copiado
        doc = Document(output_path)
        
        # Listar estilos de parágrafo disponíveis no template
        paragraph_styles = [s.name for s in doc.styles if s.type == 1]  # type 1 = paragraph
        print(f"  📋 Estilos disponíveis no template:")
        for style_name in paragraph_styles[:15]:  # Mostrar primeiros 15
            print(f"     - {style_name}")
        
        # Limpar todo o conteúdo existente (mantém cabeçalhos/rodapés)
        for _ in range(len(doc.paragraphs)):
            p = doc.paragraphs[0]
            p._element.getparent().remove(p._element)
        
        # Tentar identificar estilo para título (procurar por "Título" ou usar primeiro heading)
        title_style = None
        for style in ['Título', 'Titulo', 'Heading 1', 'Título 1']:
            if style in paragraph_styles:
                title_style = style
                break
        
        # Tentar identificar estilo para corpo de texto
        body_style = None
        for style in ['Corpo de Texto', 'Corpo de texto', 'Normal', 'Body Text']:
            if style in paragraph_styles:
                body_style = style
                break
        
        # Tentar identificar estilo para subtítulos
        subtitle_style = None
        for style in ['Título 2', 'Titulo 2', 'Heading 2', 'Subtítulo']:
            if style in paragraph_styles:
                subtitle_style = style
                break
        
        print(f"  ✓ Usando estilos: Título='{title_style}', Subtítulo='{subtitle_style}', Corpo='{body_style}'")
        
        # Adicionar título do recurso
        appeal_type = appeal_content.get('appeal_type', 'RECURSO JUDICIAL')
        if title_style:
            title = doc.add_paragraph(appeal_type.upper(), style=title_style)
        else:
            title = doc.add_paragraph()
            title_run = title.add_run(appeal_type.upper())
            title_run.font.name = 'Segoe UI'
            title_run.font.size = Pt(16)
            title_run.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Adicionar seções principais
        sections = [
            ('INTRODUÇÃO', 'introduction'),
            ('DOS FATOS', 'facts'),
            ('DOS FUNDAMENTOS', 'grounds'),
            ('DA JURISPRUDÊNCIA', 'jurisprudence'),
            ('DOS PEDIDOS', 'requests'),
            ('CONCLUSÃO', 'conclusion')
        ]
        
        for section_title, section_key in sections:
            content = appeal_content.get(section_key, '')
            if content and content.strip():
                # Adicionar título da seção usando estilo do template
                if subtitle_style:
                    doc.add_paragraph(section_title, style=subtitle_style)
                else:
                    section_para = doc.add_paragraph()
                    section_run = section_para.add_run(section_title)
                    section_run.font.name = 'Segoe UI'
                    section_run.font.size = Pt(14)
                    section_run.bold = True
                
                # Dividir por parágrafos e adicionar usando estilo do corpo do template
                paragraphs = content.split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        if body_style:
                            doc.add_paragraph(para_text.strip(), style=body_style)
                        else:
                            p = doc.add_paragraph()
                            run = p.add_run(para_text.strip())
                            run.font.name = 'Segoe UI'
                            run.font.size = Pt(11)
                # Espaço entre seções
                doc.add_paragraph()
        
        # Adicionar seções extras (se houver)
        additional_sections = appeal_content.get('additional_sections', [])
        if additional_sections:
            for section in additional_sections:
                section_title = section.get('title', 'SEÇÃO ADICIONAL')
                section_content = section.get('content', '')
                if section_content.strip():
                    # Adicionar título da seção
                    if subtitle_style:
                        doc.add_paragraph(section_title.upper(), style=subtitle_style)
                    else:
                        section_para = doc.add_paragraph()
                        section_run = section_para.add_run(section_title.upper())
                        section_run.font.name = 'Segoe UI'
                        section_run.font.size = Pt(14)
                        section_run.bold = True
                    
                    paragraphs = section_content.split('\n\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            if body_style:
                                doc.add_paragraph(para_text.strip(), style=body_style)
                            else:
                                p = doc.add_paragraph()
                                run = p.add_run(para_text.strip())
                                run.font.name = 'Segoe UI'
                                run.font.size = Pt(11)
                    doc.add_paragraph()
        
        # Salvar documento
        doc.save(output_path)
        print(f"✓ Documento criado com dados da IA: {output_path}")
        return True
        
    except Exception as e:
        print(f"✗ Erro ao criar documento DOCX: {e}")
        import traceback
        traceback.print_exc()
        # Em caso de erro, usar criação manual
        return _create_docx_manual(appeal_content, output_path)


def _create_docx_manual(appeal_content: dict, output_path: str) -> bool:
    """
    Cria documento DOCX manualmente (fallback quando template não está disponível).
    
    Args:
        appeal_content: Dicionário com o conteúdo estruturado do recurso
        output_path: Caminho onde o arquivo DOCX será salvo
        
    Returns:
        bool: True se criado com sucesso, False caso contrário
    """
    try:
        doc = Document()
        
        # Configurar margens
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)
        
        # Título do recurso
        title = doc.add_heading(appeal_content.get('appeal_type', 'RECURSO JUDICIAL').upper(), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Espaço

        # Probabilidade de êxito
        if appeal_content.get('success_probability_percent') is not None:
            doc.add_heading('PROBABILIDADE DE ÊXITO', level=1)
            percent = appeal_content.get('success_probability_percent')
            p = doc.add_paragraph(f"Estimativa: {percent}%")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(6)
            if appeal_content.get('success_probability_rationale'):
                p = doc.add_paragraph(appeal_content['success_probability_rationale'])
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(12)
        
        # Introdução
        if appeal_content.get('introduction'):
            doc.add_heading('I. INTRODUÇÃO', level=1)
            p = doc.add_paragraph(appeal_content['introduction'])
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(12)
        
        # Dos Fatos
        if appeal_content.get('facts'):
            doc.add_heading('II. DOS FATOS', level=1)
            p = doc.add_paragraph(appeal_content['facts'])
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(12)
        
        # Dos Fundamentos
        if appeal_content.get('grounds'):
            doc.add_heading('III. DOS FUNDAMENTOS', level=1)
            p = doc.add_paragraph(appeal_content['grounds'])
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(12)
        
        # Da Jurisprudência
        if appeal_content.get('jurisprudence'):
            doc.add_heading('IV. DA JURISPRUDÊNCIA', level=1)
            p = doc.add_paragraph(appeal_content['jurisprudence'])
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(12)
        
        # Seções adicionais
        section_num = 5
        for section in appeal_content.get('additional_sections', []):
            doc.add_heading(f'{section_num}. {section["title"].upper()}', level=1)
            p = doc.add_paragraph(section['content'])
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(12)
            section_num += 1
        
        # Dos Pedidos
        if appeal_content.get('requests'):
            doc.add_heading(f'{section_num}. DOS PEDIDOS', level=1)
            p = doc.add_paragraph(appeal_content['requests'])
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(12)
        
        # Conclusão
        if appeal_content.get('conclusion'):
            doc.add_heading('CONCLUSÃO', level=1)
            p = doc.add_paragraph(appeal_content['conclusion'])
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(12)
        
        # Salvar documento
        doc.save(output_path)
        print(f"✓ Documento DOCX criado: {output_path}")
        return True
        
    except Exception as e:
        print(f"✗ Erro ao criar documento DOCX: {e}")
        import traceback
        traceback.print_exc()
        return False


def generate_appeal_with_ai(
    appeal_type: str,
    sentence_analysis_dict: dict,
    user_notes: str | None = None,
    petition_path: str | None = None
) -> dict | None:
    """
    Gera um recurso judicial usando IA.
    
    Args:
        appeal_type: Tipo de recurso (Apelação, Embargos, etc)
        sentence_analysis_dict: Dicionário com a análise da sentença
        user_notes: Observações do usuário
        petition_path: Caminho da petição inicial (opcional)
        
    Returns:
        dict: Recurso gerado ou None em caso de erro
    """
    try:
        print(f"Gerando {appeal_type} com IA...")
        
        # Extrair conteúdo da petição se disponível
        petition_content = None
        if petition_path and os.path.exists(petition_path):
            from markitdown import MarkItDown
            md = MarkItDown()
            result = md.convert(petition_path)
            petition_content = result.text_content[:5000] if result.text_content else None
        
        # Gerar recurso com IA
        agent = AgentAppealGenerator(model_name="gpt-5-mini")
        appeal_result = agent.generate_appeal(
            appeal_type=appeal_type,
            sentence_analysis=sentence_analysis_dict,
            user_notes=user_notes,
            petition_content=petition_content
        )
        
        print("✓ Recurso gerado com sucesso!")
        return appeal_result
        
    except Exception as e:
        print(f"✗ Erro ao gerar recurso: {e}")
        import traceback
        traceback.print_exc()
        return None


def process_pending_appeals(batch_size: int = 10) -> int:
    """Processa um lote de recursos pendentes. Retorna quantidade processada."""
    pending_appeals = (
        JudicialAppeal.query
        .filter(JudicialAppeal.status == 'pending')
        .order_by(JudicialAppeal.created_at.asc())
        .limit(batch_size)
        .all()
    )
    
    if not pending_appeals:
        print("Nenhum recurso pendente encontrado.")
        return 0
    
    processed = 0
    
    for appeal in pending_appeals:
        try:
            # Marcar como processando
            print(f"Iniciando processamento: {appeal.id} - {appeal.appeal_type}")
            appeal.status = 'processing'
            appeal.error_message = None
            db.session.commit()
            
            # Buscar análise da sentença
            sentence = JudicialSentenceAnalysis.query.get(appeal.sentence_analysis_id)
            if not sentence or not sentence.analysis_result:
                raise Exception("Análise da sentença não disponível")
            
            # Converter análise para dict
            sentence_analysis_dict = json.loads(sentence.analysis_result)
            
            # Gerar recurso com IA
            appeal_content = generate_appeal_with_ai(
                appeal_type=appeal.appeal_type,
                sentence_analysis_dict=sentence_analysis_dict,
                user_notes=appeal.user_notes,
                petition_path=sentence.petition_file_path
            )
            
            if not appeal_content:
                raise Exception("Falha ao gerar recurso pela IA")
            
            # Salvar conteúdo como JSON
            appeal.generated_content = json.dumps(appeal_content, ensure_ascii=False, indent=2)
            
            # Criar arquivo DOCX
            upload_dir = os.path.join('uploads', 'appeals')
            os.makedirs(upload_dir, exist_ok=True)
            
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            docx_filename = f"{timestamp}_appeal_{appeal.id}_{appeal.appeal_type.replace(' ', '_')}.docx"
            docx_path = os.path.join(upload_dir, docx_filename)
            
            if create_docx_from_appeal(appeal_content, docx_path):
                appeal.generated_file_path = docx_path
            
            # Atualizar status
            appeal.processed_at = datetime.utcnow()
            appeal.status = 'completed'
            db.session.commit()
            
            processed += 1
            print(f"✓ Processado: {appeal.id} - {appeal.appeal_type}")
            
        except Exception as e:
            db.session.rollback()
            appeal.status = 'error'
            appeal.error_message = str(e)
            db.session.commit()
            import traceback
            print(f"✗ Erro ao processar {appeal.id}: {e}")
            traceback.print_exc()
    
    return processed


if __name__ == '__main__':
    with app.app_context():
        total = process_pending_appeals(batch_size=10)
        print(f"Total processado: {total}")
