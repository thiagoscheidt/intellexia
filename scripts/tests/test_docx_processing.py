"""Teste do caminho DOCX do pipeline de documentos judiciais.

Cobre as duas pontas que faziam .docx não render benefício nem resumo:

1. `DocumentProcessorService.process_document` com .docx precisa devolver
   tabelas (mesmo formato do pdfplumber), chunks e texto — antes só o PDF
   tinha tabelas, e o extrator de benefícios depende 100% de `tables`.
2. `JudicialDocumentSummaryAgent` precisa mandar TEXTO para arquivo não-PDF —
   o provider recusa .docx como anexo (400 invalid_file).

Puro Python — sem rede, sem banco, sem app Flask, sem chamada de LLM.
O .docx e o .pdf de apoio são sintéticos, criados na hora.

Rodar: uv run python scripts/tests/test_docx_processing.py
"""
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent.parent))

import fitz  # PyMuPDF
from docx import Document as DocxDocument

from app.agents.processes.judicial_document_summary_agent import JudicialDocumentSummaryAgent
from app.services.document_processor_service import DocumentProcessorService

FAILS = []


def check(label, got, expected):
    if got != expected:
        FAILS.append(f"{label}: esperado {expected!r}, obtido {got!r}")
        print(f"  ✗ {label}: esperado {expected!r}, obtido {got!r}")
    else:
        print(f"  ✓ {label}")


def check_true(label, cond, detail=""):
    if not cond:
        FAILS.append(f"{label} {detail}")
        print(f"  ✗ {label} {detail}")
    else:
        print(f"  ✓ {label}")


TMP_DIR = Path(tempfile.mkdtemp(prefix="test_docx_processing_"))

SECTION_TITLE = "5. DOENÇA NÃO RELACIONADA AO TRABALHO"
BENEFIT_A = "1234567890"
BENEFIT_B = "9876543210"


def _build_petition_docx() -> Path:
    """Petição sintética: cabeçalho, seção numerada e tabela de benefícios."""
    doc = DocxDocument()
    doc.add_paragraph("EXCELENTÍSSIMO SENHOR DOUTOR JUIZ FEDERAL DA VARA DE CHAPECÓ/SC")
    doc.add_paragraph(
        "ACME INDÚSTRIA LTDA, CNPJ 11.222.333/0001-44, vem propor AÇÃO ANULATÓRIA "
        "em face da UNIÃO FEDERAL, pelos fatos e fundamentos a seguir."
    )
    doc.add_paragraph(SECTION_TITLE)
    doc.add_paragraph(
        "Os benefícios abaixo foram indevidamente considerados no cálculo do FAP, "
        "pois a doença não guarda relação com a atividade laboral."
    )

    table = doc.add_table(rows=3, cols=5)
    header = ["VIGÊNCIA FAP", "NIT", "SEGURADO", "TIPO", "BENEFÍCIO"]
    row_a = ["2023", "12345678901", "JOAO DA SILVA", "B91", BENEFIT_A]
    row_b = ["2023", "10987654321", "MARIA SOUZA", "B31", BENEFIT_B]
    for col, value in enumerate(header):
        table.rows[0].cells[col].text = value
    for col, value in enumerate(row_a):
        table.rows[1].cells[col].text = value
    for col, value in enumerate(row_b):
        table.rows[2].cells[col].text = value

    doc.add_paragraph("Diante do exposto, requer a procedência do pedido.")

    path = TMP_DIR / "peticao_sintetica.docx"
    doc.save(str(path))
    return path


def _build_merged_header_docx() -> Path:
    """Petição com header mesclado — formato real das petições FAP em DOCX.

    "Vigências do FAP" ocupa duas colunas do grid no header, mas as linhas de
    dados usam as duas colunas separadamente. Se a célula mesclada for contada
    duas vezes, o header duplica e todas as colunas seguintes desalinham
    (CNPJ cai sob "Vigências", segurado sob "CNPJ", e assim por diante).
    """
    doc = DocxDocument()
    doc.add_paragraph("2. RESTABELECIMENTO DE BENEFÍCIO")

    table = doc.add_table(rows=2, cols=5)
    table.cell(0, 1).merge(table.cell(0, 2))
    for col, value in enumerate(["Item", "Vigências do FAP", "CNPJ", "Benefício"]):
        table.rows[0].cells[col if col < 2 else col + 1].text = value
    for col, value in enumerate(["1", "2020", "59.850.115/0001-18", BENEFIT_A, ""]):
        table.rows[1].cells[col].text = value

    path = TMP_DIR / "peticao_header_mesclado.docx"
    doc.save(str(path))
    return path


W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _first_decimal_num_id(doc) -> str:
    """numId do template cujo formato de nível 0 é decimal."""
    numbering = doc.part.numbering_part.element
    abstract_by_id = {
        node.get(f"{W_NS}abstractNumId"): node
        for node in numbering.findall(f"{W_NS}abstractNum")
    }
    for num in numbering.findall(f"{W_NS}num"):
        ref = num.find(f"{W_NS}abstractNumId")
        abstract = abstract_by_id.get(ref.get(f"{W_NS}val")) if ref is not None else None
        if abstract is None:
            continue
        for lvl in abstract.findall(f"{W_NS}lvl"):
            if lvl.get(f"{W_NS}ilvl") != "0":
                continue
            fmt = lvl.find(f"{W_NS}numFmt")
            if fmt is not None and fmt.get(f"{W_NS}val") == "decimal":
                return num.get(f"{W_NS}numId")
    raise RuntimeError("template sem numeração decimal para a fixture")


def _add_auto_numbered_heading(doc, text: str, num_id: str):
    """Título com numeração automática do Word (o número não vai no texto)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    paragraph = doc.add_paragraph(text, style="Heading 1")
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), num_id)
    num_pr.append(ilvl)
    num_pr.append(num)
    paragraph._p.get_or_add_pPr().append(num_pr)
    return paragraph


def _build_auto_numbered_docx() -> Path:
    """Petição com títulos numerados pelo Word — formato das petições reais.

    O "1."/"2." fica em numbering.xml, não no texto do parágrafo. Sem
    reconstruir esse número, nenhuma seção é detectada e a tabela perde o
    vínculo com a tese jurídica.
    """
    doc = DocxDocument()
    num_id = _first_decimal_num_id(doc)

    _add_auto_numbered_heading(doc, "ACIDENTE DE TRAJETO", num_id)
    doc.add_paragraph("Trata-se de acidente ocorrido no percurso casa-trabalho.")
    _add_auto_numbered_heading(doc, "DOENÇA NÃO RELACIONADA AO TRABALHO", num_id)
    doc.add_paragraph("Os benefícios a seguir não guardam nexo com a atividade.")

    table = doc.add_table(rows=2, cols=3)
    for col, value in enumerate(["NIT", "TIPO", "BENEFÍCIO"]):
        table.rows[0].cells[col].text = value
    for col, value in enumerate(["12345678901", "B91", BENEFIT_A]):
        table.rows[1].cells[col].text = value

    path = TMP_DIR / "peticao_numeracao_automatica.docx"
    doc.save(str(path))
    return path


def _build_small_pdf() -> Path:
    path = TMP_DIR / "peca.pdf"
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "Peça em PDF para conferir o modo de anexo.")
    pdf.save(str(path))
    pdf.close()
    return path


DOCX_PATH = _build_petition_docx()
MERGED_DOCX_PATH = _build_merged_header_docx()
AUTO_NUMBERED_DOCX_PATH = _build_auto_numbered_docx()
PDF_PATH = _build_small_pdf()


# ── 1. DocumentProcessorService.process_document com .docx ────────────────

print("\n== process_document (.docx) ==")

result = DocumentProcessorService().process_document(DOCX_PATH)

check_true(
    "extrai texto do documento",
    "ACME INDÚSTRIA LTDA" in result.full_text,
    f"-> full_text[:120]={result.full_text[:120]!r}",
)
check_true(
    "texto inclui o conteúdo da tabela (benefícios buscáveis no full_text)",
    BENEFIT_A in result.full_text and BENEFIT_B in result.full_text,
    f"-> len(full_text)={len(result.full_text)}",
)
check_true(
    "monta chunks_with_pages (contexto de seções/pedidos depende deles)",
    len(result.chunks_with_pages) > 0,
    f"-> chunks={len(result.chunks_with_pages)}",
)
check_true(
    "reporta ao menos uma página",
    result.total_pages >= 1 and len(result.pages) >= 1,
    f"-> total_pages={result.total_pages}, pages={len(result.pages)}",
)

check_true(
    "detecta a tabela de benefícios",
    len(result.tables) >= 1,
    f"-> tables={len(result.tables)}",
)

if result.tables:
    table = result.tables[0]
    rows = table.get("text") or []

    check_true(
        "tabela vem no formato do pdfplumber (page/section/text com linhas 'a | b')",
        isinstance(rows, list) and all(isinstance(row, str) for row in rows) and "page" in table,
        f"-> keys={sorted(table.keys())}, rows={rows[:1]}",
    )
    check_true(
        "primeira linha é o header",
        bool(rows) and "BENEFÍCIO" in rows[0] and "NIT" in rows[0],
        f"-> header={rows[0] if rows else None!r}",
    )
    check_true(
        "linhas de dados trazem os números de benefício",
        any(BENEFIT_A in row for row in rows) and any(BENEFIT_B in row for row in rows),
        f"-> rows={rows}",
    )
    check_true(
        "linha de dados mantém as colunas alinhadas ao header",
        any(
            row.split(" | ")[:5] == ["2023", "12345678901", "JOAO DA SILVA", "B91", BENEFIT_A]
            for row in rows
        ),
        f"-> rows={rows}",
    )
    check_true(
        "tabela carrega a seção em que aparece (mapeia a tese jurídica)",
        str(table.get("section") or "").strip() == SECTION_TITLE,
        f"-> section={table.get('section')!r}",
    )


# ── 1b. Tabela com célula mesclada no header ──────────────────────────────

print("\n== process_document (.docx com header mesclado) ==")

merged_result = DocumentProcessorService().process_document(MERGED_DOCX_PATH)
merged_rows = (merged_result.tables[0].get("text") or []) if merged_result.tables else []

check_true(
    "detecta a tabela",
    len(merged_result.tables) >= 1,
    f"-> tables={len(merged_result.tables)}",
)
check(
    "célula mesclada conta uma vez só no header",
    merged_rows[0] if merged_rows else None,
    "Item | Vigências do FAP | CNPJ | Benefício",
)
check_true(
    "linha de dados fica alinhada ao header",
    len(merged_rows) > 1 and merged_rows[1].split(" | ")[:4] == ["1", "2020", "59.850.115/0001-18", BENEFIT_A],
    f"-> row={merged_rows[1] if len(merged_rows) > 1 else None!r}",
)


# ── 1c. Títulos com numeração automática do Word ──────────────────────────

print("\n== process_document (.docx com numeração automática) ==")

auto_result = DocumentProcessorService().process_document(AUTO_NUMBERED_DOCX_PATH)
auto_sections = [page.section for page in auto_result.pages]

check_true(
    "reconstrói o número do título vindo do numbering.xml",
    any("1." in str(section or "") and "ACIDENTE DE TRAJETO" in str(section or "") for section in auto_sections)
    or any("2." in str(section or "") and "DOENÇA" in str(section or "") for section in auto_sections),
    f"-> sections={auto_sections}",
)
check_true(
    "numeração avança entre títulos (o 2º título não vira '1.')",
    any("2. DOENÇA NÃO RELACIONADA AO TRABALHO" in str(section or "") for section in auto_sections),
    f"-> sections={auto_sections}",
)
check_true(
    "tabela herda a seção do título numerado (vínculo com a tese)",
    bool(auto_result.tables)
    and "DOENÇA" in str(auto_result.tables[0].get("section") or ""),
    f"-> section={auto_result.tables[0].get('section') if auto_result.tables else None!r}",
)


# ── 2. Conteúdo enviado ao modelo pelo agente de resumo ───────────────────

print("\n== JudicialDocumentSummaryAgent: texto para não-PDF, anexo para PDF ==")

build_parts = JudicialDocumentSummaryAgent._build_document_content_parts

parts, mode = build_parts(str(DOCX_PATH), "PROMPT", document_text="TEXTO JA EXTRAIDO DA PECA")

check("docx: modo de análise é texto extraído", mode, "extracted_text")
check_true(
    "docx: nenhum anexo de arquivo é enviado (provider recusa docx)",
    all(part.get("type") != "file" for part in parts),
    f"-> types={[part.get('type') for part in parts]}",
)
check_true(
    "docx: o texto da peça vai no conteúdo",
    any("TEXTO JA EXTRAIDO DA PECA" in str(part.get("text", "")) for part in parts),
    f"-> parts={parts}",
)
check_true(
    "docx: o prompt do usuário é preservado",
    any("PROMPT" in str(part.get("text", "")) for part in parts),
    f"-> parts={parts}",
)

pdf_parts, pdf_mode = build_parts(str(PDF_PATH), "PROMPT", document_text=None)

check("pdf: modo de análise segue anexo de arquivo", pdf_mode, "file_attachment")
check_true(
    "pdf: continua mandando o arquivo como anexo (sem regressão)",
    any(part.get("type") == "file" for part in pdf_parts),
    f"-> types={[part.get('type') for part in pdf_parts]}",
)

no_text_parts, no_text_mode = build_parts(str(DOCX_PATH), "PROMPT", document_text=None)

check("docx sem texto pronto: ainda usa texto extraído", no_text_mode, "extracted_text")
check_true(
    "docx sem texto pronto: extrai o conteúdo do próprio arquivo",
    any("ACME INDÚSTRIA LTDA" in str(part.get("text", "")) for part in no_text_parts),
    f"-> parts={str(no_text_parts)[:300]}",
)

huge_text = "A" * 400_000
huge_parts, _ = build_parts(str(DOCX_PATH), "PROMPT", document_text=huge_text)
huge_chars = sum(len(str(part.get("text", ""))) for part in huge_parts)

check_true(
    "texto gigante é truncado antes de ir para o modelo",
    huge_chars < len(huge_text),
    f"-> chars enviados={huge_chars}",
)


# ── resultado ─────────────────────────────────────────────────────────────

print()
if FAILS:
    print(f"FALHOU: {len(FAILS)} check(s) com problema")
    for fail in FAILS:
        print(f"  - {fail}")
    sys.exit(1)

print("OK: todos os checks passaram")
