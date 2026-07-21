"""
Teste dos marcadores de imagem na extração de DOCX do Revisor FAP.

Regressão: petições trazem provas como prints embutidos (telas do FAP, CATs).
A extração de texto descartava as imagens silenciosamente e a IA apontava
"documento em falta" para provas que estavam no documento — como imagem.

Uso: uv run python tests/test_fap_review_docx_images.py
"""

import base64
import io
import os
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
os.environ.setdefault('OPENAI_API_KEY', 'test-key')

from docx import Document  # noqa: E402

from app.blueprints.fap_review import _extract_text_from_document, IMAGE_MARKER  # noqa: E402

# PNG 1x1 transparente
PNG_1X1 = base64.b64decode(
    'iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNkYPhfDwAChwGA60e6kgAAAABJRU5ErkJggg=='
)

PASSED = 0
FAILED = 0


def check(label: str, condition: bool, detail: str = "") -> None:
    global PASSED, FAILED
    if condition:
        PASSED += 1
        print(f"  ✓ {label}")
    else:
        FAILED += 1
        print(f"  ✗ {label} {detail}")


def run():
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / 'peticao.docx'
        doc = Document()
        doc.add_paragraph('Comunicação de Acidente de Trabalho (CAT)')
        doc.add_picture(io.BytesIO(PNG_1X1))
        doc.add_paragraph('DOS PEDIDOS')
        doc.save(str(path))

        text = _extract_text_from_document(str(path))
        lines = text.split('\n')

    print("[1] Imagem embutida vira marcador no texto extraído")
    check("marcador presente", IMAGE_MARKER in text)
    cat_idx = lines.index('Comunicação de Acidente de Trabalho (CAT)')
    marker_idx = next(i for i, l in enumerate(lines) if IMAGE_MARKER in l)
    pedidos_idx = lines.index('DOS PEDIDOS')
    check("marcador na posição correta (entre CAT e PEDIDOS)", cat_idx < marker_idx < pedidos_idx,
          f"(cat={cat_idx}, marcador={marker_idx}, pedidos={pedidos_idx})")

    print("[1b] Imagem dentro de célula de tabela também vira marcador")
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / 'tabela.docx'
        doc = Document()
        doc.add_paragraph('Telas do FAP')
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = 'Vigência 2024'
        table.rows[0].cells[1].paragraphs[0].add_run().add_picture(io.BytesIO(PNG_1X1))
        doc.save(str(path))
        text = _extract_text_from_document(str(path))
    check("marcador presente na linha da tabela", IMAGE_MARKER in text)
    table_line = next((l for l in text.split('\n') if 'Vigência 2024' in l), '')
    check("marcador na mesma linha da célula", IMAGE_MARKER in table_line, f"(linha: {table_line!r})")

    print("[2] DOCX sem imagens não ganha marcador")
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / 'simples.docx'
        doc = Document()
        doc.add_paragraph('Texto sem imagens')
        doc.save(str(path))
        text = _extract_text_from_document(str(path))
    check("sem marcador", IMAGE_MARKER not in text)


if __name__ == '__main__':
    run()
    print(f"\nResultado: {PASSED} ok, {FAILED} falhas")
    sys.exit(1 if FAILED else 0)
