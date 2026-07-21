"""
Teste do preview HTML de documentos DOCX no Revisor de Petições FAP.

Verifica que o helper `render_docx_preview_html` converte um DOCX real em HTML
preservando parágrafos, títulos, negrito e tabelas — base do modal "Ver Página"
para documentos que o navegador não renderiza nativamente.

Uso: uv run python tests/test_fap_review_docx_preview.py
"""

import os
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
os.environ.setdefault('OPENAI_API_KEY', 'test-key')

from docx import Document  # noqa: E402

from app.utils.document_utils import render_docx_preview_html  # noqa: E402

PASSED = 0
FAILED = 0


def check(label: str, condition: bool, detail: str = "") -> None:
    global PASSED, FAILED
    if condition:
        PASSED += 1
        print(f"  ✓ {label}")
    else:
        FAILED += 1
        print(f"  ✗ {label} {detail}")


def build_sample_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading('SÍNTESE FÁTICA', level=1)
    paragraph = doc.add_paragraph('A autora MARFRIG GLOBAL FOODS S.A está sujeita ao recolhimento ')
    paragraph.add_run('da contribuição ao SAT').bold = True
    paragraph.add_run(' sobre a folha de pagamento.')
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = 'NIT'
    table.rows[0].cells[1].text = '123.45678.90-1'
    doc.save(str(path))


def test_render_docx_preview_html():
    print("[1] Conversão DOCX → HTML")
    with tempfile.TemporaryDirectory() as tmp:
        docx_path = Path(tmp) / 'peticao.docx'
        build_sample_docx(docx_path)
        html = render_docx_preview_html(docx_path)

    check("retorna string não vazia", isinstance(html, str) and bool(html.strip()))
    check("preserva título como heading", '<h1>' in html and 'SÍNTESE FÁTICA' in html)
    check("preserva negrito", '<strong>' in html and 'contribuição ao SAT' in html)
    check("preserva tabela", '<table>' in html and '123.45678.90-1' in html)
    check("texto do parágrafo presente", 'MARFRIG GLOBAL FOODS S.A' in html)


def test_render_docx_preview_html_invalid_file():
    print("[2] Arquivo inválido gera ValueError")
    with tempfile.TemporaryDirectory() as tmp:
        bad_path = Path(tmp) / 'nao_e_docx.docx'
        bad_path.write_bytes(b'conteudo que nao e um docx')
        try:
            render_docx_preview_html(bad_path)
            check("levantou ValueError", False, "(não levantou)")
        except ValueError:
            check("levantou ValueError", True)


if __name__ == '__main__':
    test_render_docx_preview_html()
    test_render_docx_preview_html_invalid_file()
    print(f"\nResultado: {PASSED} ok, {FAILED} falhas")
    sys.exit(1 if FAILED else 0)
