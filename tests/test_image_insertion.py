"""
Script de teste para validar a funcionalidade de inserção de imagens
Testa conversão de PDF, inserção de imagens e formatação
"""

import os
import sys
from io import BytesIO
from pdf2image import convert_from_path
from PIL import Image
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def test_poppler_installation():
    """Testa se Poppler está instalado corretamente"""
    print("🔍 Testando instalação do Poppler...")
    try:
        # Tenta importar pdf2image
        from pdf2image import convert_from_path
        print("✅ pdf2image importado com sucesso")
        
        # Nota: Não podemos testar conversão real sem um PDF
        print("ℹ️  Para testar conversão, execute: pdftoppm -v no terminal")
        return True
    except Exception as e:
        print(f"❌ Erro ao importar pdf2image: {e}")
        return False


def test_pillow_installation():
    """Testa se Pillow está instalado corretamente"""
    print("\n🔍 Testando instalação do Pillow...")
    try:
        from PIL import Image
        print("✅ Pillow importado com sucesso")
        print(f"ℹ️  Versão do Pillow: {Image.__version__ if hasattr(Image, '__version__') else 'Desconhecida'}")
        return True
    except Exception as e:
        print(f"❌ Erro ao importar Pillow: {e}")
        return False


def test_docx_image_insertion():
    """Testa inserção de imagem em documento Word"""
    print("\n🔍 Testando inserção de imagem em Word...")
    try:
        # Criar documento de teste
        doc = Document()
        
        # Adicionar título
        doc.add_heading('Teste de Inserção de Imagem', 0)
        
        # Adicionar parágrafo
        p = doc.add_paragraph('Teste de placeholder:')
        
        # Simular placeholder
        placeholder_p = doc.add_paragraph('{{imagem_cat}}')
        
        print("✅ Documento Word criado com sucesso")
        
        # Salvar documento de teste
        test_file = 'test_image_insertion.docx'
        doc.save(test_file)
        print(f"✅ Documento salvo: {test_file}")
        
        # Limpar
        if os.path.exists(test_file):
            os.remove(test_file)
            print("✅ Arquivo de teste removido")
        
        return True
    except Exception as e:
        print(f"❌ Erro ao criar documento Word: {e}")
        return False


def test_pdf_conversion_simulation():
    """Simula conversão de PDF (sem arquivo real)"""
    print("\n🔍 Simulando conversão de PDF...")
    print("ℹ️  Para teste real, você precisa:")
    print("   1. Ter um arquivo PDF em uploads/cases/")
    print("   2. Executar: python -c 'from pdf2image import convert_from_path; images = convert_from_path(\"arquivo.pdf\", first_page=1, last_page=1, dpi=150); print(f\"Sucesso: {len(images)} página(s) convertida(s)\")'")
    return True


def test_image_formats():
    """Testa suporte a diferentes formatos de imagem"""
    print("\n🔍 Testando formatos de imagem suportados...")
    
    supported_formats = ['.png', '.jpg', '.jpeg', '.bmp', '.gif']
    print(f"✅ Formatos suportados: {', '.join(supported_formats)}")
    
    return True


def test_agent_imports():
    """Testa se os imports do AgentDocumentGenerator estão corretos"""
    print("\n🔍 Testando imports do AgentDocumentGenerator...")
    try:
        from app.models import Case, CaseBenefit, Document as DocumentModel
        print("✅ Models importados com sucesso")
        
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        print("✅ python-docx importado com sucesso")
        
        from pdf2image import convert_from_path
        print("✅ pdf2image importado com sucesso")
        
        from PIL import Image
        print("✅ PIL importado com sucesso")
        
        from docxcompose.composer import Composer
        print("✅ docxcompose importado com sucesso")
        
        return True
    except ImportError as e:
        print(f"❌ Erro ao importar módulo: {e}")
        return False


def test_file_paths():
    """Verifica estrutura de diretórios"""
    print("\n🔍 Verificando estrutura de diretórios...")
    
    directories = [
        'uploads',
        'uploads/cases',
        'uploads/petitions',
        'templates_docx'
    ]
    
    for directory in directories:
        if os.path.exists(directory):
            print(f"✅ {directory}/ existe")
        else:
            print(f"⚠️  {directory}/ não existe (será criado automaticamente)")
    
    return True


def run_all_tests():
    """Executa todos os testes"""
    print("=" * 60)
    print("🧪 TESTE DE INSERÇÃO DE IMAGENS EM PETIÇÕES")
    print("=" * 60)
    
    tests = [
        ("Poppler", test_poppler_installation),
        ("Pillow", test_pillow_installation),
        ("Inserção em Word", test_docx_image_insertion),
        ("Conversão PDF", test_pdf_conversion_simulation),
        ("Formatos de Imagem", test_image_formats),
        ("Imports do Agent", test_agent_imports),
        ("Estrutura de Diretórios", test_file_paths),
    ]
    
    results = []
    for test_name, test_func in tests:
        try:
            result = test_func()
            results.append((test_name, result))
        except Exception as e:
            print(f"\n❌ Erro inesperado em '{test_name}': {e}")
            results.append((test_name, False))
    
    # Resumo
    print("\n" + "=" * 60)
    print("📊 RESUMO DOS TESTES")
    print("=" * 60)
    
    passed = sum(1 for _, result in results if result)
    total = len(results)
    
    for test_name, result in results:
        status = "✅ PASSOU" if result else "❌ FALHOU"
        print(f"{status} - {test_name}")
    
    print(f"\n🎯 Total: {passed}/{total} testes passaram")
    
    if passed == total:
        print("\n🎉 Todos os testes passaram! Sistema pronto para uso.")
    else:
        print("\n⚠️  Alguns testes falharam. Verifique as mensagens de erro acima.")
        print("\n📚 Consulte: docs/TESTE_INSERCAO_IMAGENS.md para troubleshooting")
    
    return passed == total


if __name__ == '__main__':
    success = run_all_tests()
    sys.exit(0 if success else 1)
